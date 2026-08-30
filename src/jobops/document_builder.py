from __future__ import annotations

import os
import re
import subprocess
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from pathlib import PurePosixPath
from typing import Any
from xml.etree import ElementTree as ET

from lxml import etree as LET

from .application_narrative import validate_application_narrative_text
from .claims import external_use_decision
from .errors import JobOpsError
from .util import canonical_json, sha256_bytes, sha256_file, stable_id


NAVY = "12335B"
BLUE = "285F9B"
GRAY = "606A78"
LIGHT = "E8EEF5"
BLACK = "111111"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CORE_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
CUSTOM_PROPERTIES_NS = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
VT_NS = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
MAX_DOCX_PACKAGE_MEMBERS = 10_000
MAX_DOCX_PACKAGE_UNCOMPRESSED_BYTES = 256 * 1024 * 1024
MAX_DOCX_PACKAGE_PART_BYTES = 64 * 1024 * 1024
MAX_DOCX_PACKAGE_COMPRESSION_RATIO = 200
ALLOWED_TEMPLATE_SLOTS = {
    "CANDIDATE_NAME", "TARGET_ROLE", "SUMMARY", "EXPERIENCE_BULLET",
    "PROJECT", "SKILLS", "EDUCATION", "COVER_LETTER", "APPLICATION_NARRATIVE",
}


@dataclass(frozen=True)
class CoverLetterNarrative:
    candidate_display_name: str
    company: str
    target_role: str
    paragraphs: tuple[str, ...]
    text: str
    claim_ids: tuple[str, ...]


def _empty_custom_properties() -> bytes:
    ET.register_namespace("", CUSTOM_PROPERTIES_NS)
    ET.register_namespace("vt", VT_NS)
    return ET.tostring(ET.Element(f"{{{CUSTOM_PROPERTIES_NS}}}Properties"), encoding="utf-8", xml_declaration=True)


def _parse_xml_preserving_namespaces(data: bytes) -> Any:
    """Parse OOXML without renaming prefixes referenced by mc:Ignorable."""

    try:
        parser = LET.XMLParser(resolve_entities=False, no_network=True, recover=False, huge_tree=False)
        return LET.fromstring(data, parser=parser)
    except LET.XMLSyntaxError as exc:
        raise JobOpsError("DOCX_XML_INVALID", "An OOXML document part could not be parsed safely.") from exc


def _serialize_xml_preserving_namespaces(root: Any) -> bytes:
    return LET.tostring(root, encoding="UTF-8", xml_declaration=True, standalone=True)


@dataclass(frozen=True)
class TemplateFingerprint:
    master_sha256: str
    page_geometry: tuple[tuple[int, int, int, int, int, int], ...]
    style_ids: tuple[str, ...]
    table_grids: tuple[tuple[int, ...], ...]
    headers: tuple[str, ...]
    footers: tuple[str, ...]
    hyperlinks: tuple[str, ...]
    package_parts: tuple[tuple[str, int, str], ...]

    def as_dict(self) -> dict[str, object]:
        return {
            **self.__dict__,
            "page_geometry": [list(item) for item in self.page_geometry],
            "table_grids": [list(item) for item in self.table_grids],
            "package_parts": [list(item) for item in self.package_parts],
        }


def _zip_inventory(path: Path) -> tuple[tuple[str, int, str], ...]:
    try:
        with zipfile.ZipFile(path) as archive:
            members = archive.infolist()
            if not members or len(members) > MAX_DOCX_PACKAGE_MEMBERS:
                raise JobOpsError("DOCX_PACKAGE_UNSAFE", "The DOCX package contains an invalid number of parts.")
            names: set[str] = set()
            total = 0
            document_parts = 0
            for info in members:
                normalized = PurePosixPath(info.filename.replace("\\", "/"))
                folded = normalized.as_posix().casefold()
                if (
                    normalized.is_absolute()
                    or ".." in normalized.parts
                    or "\\" in info.filename
                    or not normalized.parts
                    or folded in names
                ):
                    raise JobOpsError("DOCX_PACKAGE_UNSAFE", "The DOCX package contains an unsafe or duplicate part name.")
                names.add(folded)
                if info.is_dir():
                    continue
                if info.flag_bits & 0x1:
                    raise JobOpsError("DOCX_PACKAGE_ENCRYPTED", "Encrypted DOCX package parts are not supported.")
                if info.file_size < 0 or info.file_size > MAX_DOCX_PACKAGE_PART_BYTES:
                    raise JobOpsError("DOCX_PACKAGE_UNSAFE", "A DOCX package part exceeds the bounded size limit.")
                total += int(info.file_size)
                if total > MAX_DOCX_PACKAGE_UNCOMPRESSED_BYTES:
                    raise JobOpsError("DOCX_PACKAGE_UNSAFE", "The DOCX package exceeds the bounded expanded-size limit.")
                if int(info.file_size) / max(1, int(info.compress_size)) > MAX_DOCX_PACKAGE_COMPRESSION_RATIO:
                    raise JobOpsError("DOCX_PACKAGE_COMPRESSION_UNSAFE", "A DOCX package part has an unsafe compression ratio.")
                if info.filename == "word/document.xml":
                    document_parts += 1
            if document_parts != 1 or "[content_types].xml" not in names:
                raise JobOpsError("DOCX_PACKAGE_INVALID", "The DOCX package is missing its unique main document structure.")
            return tuple(
                sorted((info.filename, info.file_size, sha256_bytes(archive.read(info))) for info in members if not info.is_dir())
            )
    except JobOpsError:
        raise
    except (OSError, RuntimeError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise JobOpsError("DOCX_PACKAGE_INVALID", "The DOCX package could not be read safely.") from exc


def template_fingerprint(path: Path) -> TemplateFingerprint:
    """Return a structural fingerprint without persisting resume content."""
    from docx import Document
    from docx.oxml.ns import qn

    package_inventory = _zip_inventory(path)
    try:
        document = Document(str(path))
    except Exception as exc:
        raise JobOpsError("DOCX_PACKAGE_INVALID", "The validated DOCX package could not be opened as an editable document.") from exc
    geometry = tuple(
        (
            int(section.page_width or 0), int(section.page_height or 0),
            int(section.top_margin or 0), int(section.bottom_margin or 0),
            int(section.left_margin or 0), int(section.right_margin or 0),
        )
        for section in document.sections
    )
    styles = tuple(sorted(style.style_id for style in document.styles if style.style_id))
    grids: list[tuple[int, ...]] = []
    for table in document.tables:
        grid = table._tbl.tblGrid
        grids.append(tuple(int(col.get(qn("w:w")) or 0) for col in grid.gridCol_lst) if grid is not None else tuple())
    headers = tuple(sha256_bytes("\n".join(p.text for p in section.header.paragraphs).encode("utf-8")) for section in document.sections)
    footers = tuple(sha256_bytes("\n".join(p.text for p in section.footer.paragraphs).encode("utf-8")) for section in document.sections)
    hyperlinks: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/_rels/") and name.endswith(".rels"):
                root = ET.fromstring(archive.read(name))
                for rel in root:
                    if rel.attrib.get("TargetMode") == "External":
                        hyperlinks.append(rel.attrib.get("Target", ""))
    return TemplateFingerprint(
        master_sha256=sha256_file(path), page_geometry=geometry, style_ids=styles,
        table_grids=tuple(grids), headers=headers, footers=footers,
        hyperlinks=tuple(sorted(hyperlinks)), package_parts=package_inventory,
    )


def discover_template_slots(path: Path) -> list[str]:
    """Return only literal slots the preserve-only tailor can actually replace."""

    _zip_inventory(path)
    found: set[str] = set()
    try:
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                if not (
                    name == "word/document.xml"
                    or name.startswith("word/header")
                    or name.startswith("word/footer")
                ) or not name.endswith(".xml"):
                    continue
                data = archive.read(name)
                for raw in re.findall(rb"\{\{([A-Z_]{2,40})\}\}", data):
                    slot = raw.decode("ascii")
                    if slot in ALLOWED_TEMPLATE_SLOTS:
                        found.add(slot)
    except (OSError, RuntimeError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise JobOpsError("DOCX_PACKAGE_INVALID", "The DOCX template slots could not be inspected safely.") from exc
    return sorted(found)


def inspect_docx_text_blocks(path: Path) -> list[dict[str, Any]]:
    """Inspect stable body-paragraph positions without persisting their text.

    The returned text is intended only for an authenticated localhost review response.
    A persisted manifest stores the hashes and opaque block references instead.
    """

    _zip_inventory(path)
    master_hash = sha256_file(path)
    try:
        with zipfile.ZipFile(path) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))
    except (OSError, RuntimeError, KeyError, ET.ParseError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise JobOpsError("DOCX_PACKAGE_INVALID", "The editable body blocks could not be inspected safely.") from exc
    blocks: list[dict[str, Any]] = []
    line_number = 0
    paragraphs = [node for node in root.iter() if node.tag == f"{{{W_NS}}}p"]
    for paragraph_index, paragraph in enumerate(paragraphs):
        text = "".join(node.text or "" for node in paragraph.iter() if node.tag == f"{{{W_NS}}}t").strip()
        if not text:
            continue
        line_number += 1
        text_hash = sha256_bytes(text.encode("utf-8"))
        style = next((node.attrib.get(f"{{{W_NS}}}val", "") for node in paragraph.iter() if node.tag == f"{{{W_NS}}}pStyle"), "")
        is_list = any(node.tag == f"{{{W_NS}}}numPr" for node in paragraph.iter()) or bool(re.match(r"^[\u2022\u25cf\u25aa\uf0b7\-]\s*", text))
        blocks.append({
            "block_ref": stable_id("RBL", master_hash, "word/document.xml", str(paragraph_index), text_hash),
            "part_name": "word/document.xml", "paragraph_index": paragraph_index,
            "line_number": line_number, "text": text, "text_sha256": text_hash,
            "text_length": len(text), "style_id": style, "is_list": is_list,
        })
    return blocks


def tailor_master_resume_with_manifest(
    master_path: Path,
    output_path: Path,
    *,
    manifest: dict[str, Any],
    replacements: list[dict[str, str]],
    external_claim_set: dict[str, Any],
    synthetic: bool = False,
) -> dict[str, object]:
    """Replace approved body blocks with exact externally approved Claim wording."""

    from .external_claims import validate_external_claim_set_integrity
    from .resume_tailoring import validate_resume_tailoring_manifest_integrity

    if master_path.resolve() == output_path.resolve():
        raise JobOpsError("MASTER_OVERWRITE_FORBIDDEN", "Tailoring must write a new copy, never the retained master resume.")
    if master_path.suffix.casefold() != ".docx":
        raise JobOpsError("DOCX_MASTER_REQUIRED", "Manifest tailoring requires an editable DOCX master.")
    validate_resume_tailoring_manifest_integrity(manifest)
    validate_external_claim_set_integrity(external_claim_set)
    before = template_fingerprint(master_path)
    if manifest.get("master_resume_sha256") != before.master_sha256:
        raise JobOpsError("TAILORING_MANIFEST_STALE", "The approved tailoring map belongs to a different Master Resume.")
    inspected = {item["block_ref"]: item for item in inspect_docx_text_blocks(master_path)}
    approved_blocks = {item["block_ref"]: item for item in manifest.get("blocks", [])}
    approved_claims = {item["claim_id"]: item for item in external_claim_set.get("claims", [])}
    normalized: list[dict[str, Any]] = []
    seen_blocks: set[str] = set()
    for item in replacements:
        block_ref, claim_id = str(item.get("block_ref", "")), str(item.get("claim_id", ""))
        if block_ref in seen_blocks or block_ref not in approved_blocks or block_ref not in inspected:
            raise JobOpsError("TAILORING_REPLACEMENT_INVALID", "A requested resume block is missing, duplicated, or not approved.")
        seen_blocks.add(block_ref)
        approved_block, current_block = approved_blocks[block_ref], inspected[block_ref]
        if approved_block.get("original_text_sha256") != current_block.get("text_sha256"):
            raise JobOpsError("TAILORING_MANIFEST_STALE", "An approved resume block changed after review.")
        claim = approved_claims.get(claim_id)
        if claim is None or claim.get("approved_for_external") is not True or "resume" not in claim.get("allowed_uses", []):
            raise JobOpsError("TAILORING_CLAIM_NOT_APPROVED", "A selected Claim is not currently approved for resume use.")
        if claim.get("category") != approved_block.get("category"):
            raise JobOpsError("TAILORING_CATEGORY_MISMATCH", "A selected Claim does not match the approved resume-block category.")
        wording = str((claim.get("allowed_wording") or [""])[0]).strip()
        if not wording or len(wording) > int(approved_block.get("maximum_characters", 0)):
            raise JobOpsError("TAILORING_CONTENT_TOO_LONG", "Approved Claim wording exceeds this block's safe review length.")
        if not synthetic and any(marker in wording.casefold() for marker in ("jobops", "secure-ref:", "evidence-gated")):
            raise JobOpsError("INTERNAL_MARKER_FORBIDDEN", "Application material cannot contain internal workflow markers.")
        normalized.append({**current_block, "claim_id": claim_id, "replacement": wording})
    if not normalized:
        raise JobOpsError("TAILORING_REPLACEMENT_EMPTY", "At least one approved resume block must be tailored.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_output = output_path.with_name(f".{output_path.name}.jobflow-{uuid.uuid4().hex}.tmp")
    changed_parts: set[str] = set()
    changes: list[dict[str, Any]] = []
    try:
        with zipfile.ZipFile(master_path, "r") as source, zipfile.ZipFile(temporary_output, "x") as target:
            for info in source.infolist():
                name = info.filename
                data = source.read(name)
                if name == "docProps/custom.xml" and not synthetic:
                    data = _empty_custom_properties()
                    changed_parts.add(name)
                if name == "word/document.xml":
                    root = _parse_xml_preserving_namespaces(data)
                    paragraphs = [node for node in root.iter() if node.tag == f"{{{W_NS}}}p"]
                    for replacement in normalized:
                        index = int(replacement["paragraph_index"])
                        if not 0 <= index < len(paragraphs):
                            raise JobOpsError("TAILORING_MANIFEST_STALE", "An approved resume block no longer exists.")
                        text_nodes = [node for node in paragraphs[index].iter() if node.tag == f"{{{W_NS}}}t"]
                        if not text_nodes:
                            raise JobOpsError("TAILORING_MANIFEST_STALE", "An approved resume block no longer contains editable text.")
                        text_nodes[0].text = replacement["replacement"]
                        for node in text_nodes[1:]:
                            node.text = ""
                        changes.append({
                            "block_ref": replacement["block_ref"], "claim_id": replacement["claim_id"],
                            "replacement_sha256": sha256_bytes(replacement["replacement"].encode("utf-8")),
                        })
                    data = _serialize_xml_preserving_namespaces(root)
                    changed_parts.add(name)
                if name == "docProps/core.xml" and not synthetic:
                    root = _parse_xml_preserving_namespaces(data)
                    for tag in (f"{{{DC_NS}}}creator", f"{{{CP_NS}}}lastModifiedBy", f"{{{DC_NS}}}description"):
                        node = root.find(tag)
                        if node is not None:
                            node.text = ""
                    data = _serialize_xml_preserving_namespaces(root)
                    changed_parts.add(name)
                target.writestr(info, data)
    except Exception:
        temporary_output.unlink(missing_ok=True)
        raise
    if sha256_file(master_path) != before.master_sha256:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("MASTER_CHANGED_DURING_TAILORING", "The retained master changed while its copy was tailored.")
    try:
        after = template_fingerprint(temporary_output)
    except Exception:
        temporary_output.unlink(missing_ok=True)
        raise
    if after.page_geometry != before.page_geometry or after.style_ids != before.style_ids or after.table_grids != before.table_grids or after.hyperlinks != before.hyperlinks:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("TEMPLATE_STRUCTURE_CHANGED", "The tailored copy changed page geometry, styles, tables, or hyperlinks.")
    before_parts = {name: (size, digest) for name, size, digest in before.package_parts}
    after_parts = {name: (size, digest) for name, size, digest in after.package_parts}
    expected_changes = changed_parts | ({"docProps/custom.xml"} if "docProps/custom.xml" in before_parts and not synthetic else set())
    unexplained = sorted(name for name in set(before_parts) | set(after_parts) if before_parts.get(name) != after_parts.get(name) and name not in expected_changes)
    if unexplained:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("UNEXPLAINED_PACKAGE_CHANGE", "An unrelated DOCX package part changed.", parts=unexplained)
    diff: dict[str, Any] = {
        "master_sha256": before.master_sha256, "output_sha256": sha256_file(temporary_output),
        "manifest_content_hash": manifest.get("content_hash"), "claim_set_content_hash": external_claim_set.get("content_hash"),
        "changed_parts": sorted(changed_parts), "block_changes": changes,
        "preserved": {"page_geometry": True, "styles": True, "table_grids": True, "hyperlinks": True, "unexplained_package_changes": []},
    }
    diff["diff_sha256"] = sha256_bytes(canonical_json(diff))
    try:
        os.replace(temporary_output, output_path)
    except OSError as exc:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("DOCUMENT_OUTPUT_COMMIT_FAILED", "The validated tailored document could not be committed atomically.") from exc
    return diff


def _claim_wordings(claims: list[dict[str, Any]]) -> set[str]:
    return {wording for _, wording, _ in _assert_claims(claims)}


def tailor_master_resume(
    master_path: Path,
    output_path: Path,
    *,
    replacements: dict[str, str],
    claims: list[dict[str, Any]] | None = None,
    external_claim_set: dict[str, Any] | None = None,
    synthetic: bool = False,
) -> dict[str, object]:
    """Patch explicit template slots in a copy while preserving unrelated OOXML parts.

    Slots are literal ``{{NAME}}`` text nodes in body, table, header, or footer XML.
    Claim-bearing slots are exact-wording gated. The retained master is never opened
    for writing and its hash is checked again after the output is produced.
    """
    if master_path.resolve() == output_path.resolve():
        raise JobOpsError("MASTER_OVERWRITE_FORBIDDEN", "Tailoring must write a new copy, never the retained master resume.")
    if master_path.suffix.casefold() != ".docx":
        raise JobOpsError("DOCX_MASTER_REQUIRED", "PDF masters may be securely retained, but editable tailoring requires a DOCX master.")
    before = template_fingerprint(master_path)
    unknown = sorted(set(replacements) - ALLOWED_TEMPLATE_SLOTS)
    if unknown:
        raise JobOpsError("UNKNOWN_TEMPLATE_SLOT", "Template replacement contains an unsupported slot.", slots=unknown)
    gated = {"SUMMARY", "EXPERIENCE_BULLET", "PROJECT", "SKILLS", "EDUCATION", "COVER_LETTER", "APPLICATION_NARRATIVE"}
    if (claims is None) == (external_claim_set is None):
        raise JobOpsError("MATERIAL_CLAIM_SOURCE_INVALID", "Choose exactly one approved Claim source for resume tailoring.")
    wordings = (
        _claim_wordings(claims or [])
        if claims is not None
        else {str(item["allowed_wording"][0]) for item in _assert_external_claims(external_claim_set or {}, use="resume")}
    )
    unsupported = sorted(key for key in replacements if key in gated and replacements[key] not in wordings)
    if unsupported:
        raise JobOpsError("TEMPLATE_CONTENT_NOT_CLAIM_GATED", "Every tailored factual slot must use exact current approved Claim wording.", slots=unsupported)
    forbidden = ("JobOps", "evidence-gated", "secure-ref:")
    if not synthetic and any(token.casefold() in value.casefold() for value in replacements.values() for token in forbidden):
        raise JobOpsError("INTERNAL_MARKER_FORBIDDEN", "Real-mode material cannot contain internal workflow markers.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_output = output_path.with_name(f".{output_path.name}.jobflow-{uuid.uuid4().hex}.tmp")
    replaced: list[dict[str, str]] = []
    changed_parts: set[str] = set()
    try:
        with zipfile.ZipFile(master_path, "r") as source, zipfile.ZipFile(temporary_output, "x") as target:
            for info in source.infolist():
                name = info.filename
                data = source.read(name)
                if name == "docProps/custom.xml" and not synthetic:
                    data = _empty_custom_properties()
                    changed_parts.add(name)
                if (name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer")) and name.endswith(".xml"):
                    original = data
                    for slot, value in replacements.items():
                        marker = "{{" + slot + "}}"
                        count = data.count(marker.encode("utf-8"))
                        if count:
                            data = data.replace(marker.encode("utf-8"), value.encode("utf-8"))
                            replaced.append({"slot": slot, "part": name, "occurrences": str(count), "replacement_sha256": sha256_bytes(value.encode("utf-8"))})
                    if data != original:
                        changed_parts.add(name)
                if name == "docProps/core.xml" and not synthetic:
                    root = _parse_xml_preserving_namespaces(data)
                    for tag in (f"{{{DC_NS}}}creator", f"{{{CP_NS}}}lastModifiedBy", f"{{{DC_NS}}}description"):
                        node = root.find(tag)
                        if node is not None:
                            node.text = ""
                    data = _serialize_xml_preserving_namespaces(root)
                    changed_parts.add(name)
                target.writestr(info, data)
    except Exception:
        temporary_output.unlink(missing_ok=True)
        raise
    missing = sorted(set(replacements) - {item["slot"] for item in replaced})
    if missing:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("TEMPLATE_SLOT_NOT_FOUND", "Every requested slot must exist as a stable literal template marker.", slots=missing)
    if sha256_file(master_path) != before.master_sha256:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("MASTER_CHANGED_DURING_TAILORING", "The retained master resume changed while the copy was being tailored.")
    try:
        after = template_fingerprint(temporary_output)
    except Exception:
        temporary_output.unlink(missing_ok=True)
        raise
    if after.page_geometry != before.page_geometry or after.style_ids != before.style_ids or after.table_grids != before.table_grids or after.hyperlinks != before.hyperlinks:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("TEMPLATE_STRUCTURE_CHANGED", "Page geometry, styles, tables, or hyperlinks changed outside the approved text slots.")
    before_parts = {name: (size, digest) for name, size, digest in before.package_parts}
    after_parts = {name: (size, digest) for name, size, digest in after.package_parts}
    expected_changes = changed_parts | ({"docProps/custom.xml"} if "docProps/custom.xml" in before_parts and not synthetic else set())
    unexplained = sorted(
        name for name in set(before_parts) | set(after_parts)
        if before_parts.get(name) != after_parts.get(name) and name not in expected_changes
    )
    if unexplained:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("UNEXPLAINED_PACKAGE_CHANGE", "A preserve-only package part changed.", parts=unexplained)
    diff = {
        "master_sha256": before.master_sha256,
        "output_sha256": sha256_file(temporary_output),
        "changed_parts": sorted(changed_parts),
        "removed_metadata_parts": sorted(expected_changes - changed_parts),
        "slot_changes": replaced,
        "preserved": {
            "page_geometry": True, "styles": True, "table_grids": True,
            "headers_footers_except_explicit_slots": True, "hyperlinks": True,
            "unexplained_package_changes": [],
        },
    }
    diff["diff_sha256"] = sha256_bytes(canonical_json(diff))
    try:
        os.replace(temporary_output, output_path)
    except OSError as exc:
        temporary_output.unlink(missing_ok=True)
        raise JobOpsError("DOCUMENT_OUTPUT_COMMIT_FAILED", "The validated tailored document could not be committed atomically.") from exc
    return diff


def _modules():
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    return Document, WD_SECTION, WD_STYLE_TYPE, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor


def _set_run(run, *, size: float = 11, bold: bool = False, color: str = BLACK, italic: bool = False) -> None:
    _, _, _, _, _, qn, _, Pt, RGBColor = _modules()
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _configure_document(document) -> None:
    _, _, WD_STYLE_TYPE, WD_ALIGN_PARAGRAPH, _, qn, Inches, Pt, RGBColor = _modules()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, NAVY),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name, size, bold, color, alignment, after in (
        ("JobOps Title", 20, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, 2),
        ("JobOps Subtitle", 11.5, True, GRAY, WD_ALIGN_PARAGRAPH.CENTER, 10),
        ("JobOps Letter Title", 18, True, NAVY, WD_ALIGN_PARAGRAPH.LEFT, 4),
        ("JobOps Letter Subtitle", 11, True, GRAY, WD_ALIGN_PARAGRAPH.LEFT, 16),
    ):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    document.core_properties.author = "JobOps"
    document.core_properties.last_modified_by = "JobOps"
    document.core_properties.comments = "Generated from approved synthetic claims during JobOps validation."


def _save_document_atomic(document: Any, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.jobflow-{uuid.uuid4().hex}.tmp")
    try:
        document.save(str(temporary))
        if not temporary.is_file() or temporary.stat().st_size < 1:
            raise OSError("document writer produced no output")
        os.replace(temporary, path)
    except Exception as exc:
        temporary.unlink(missing_ok=True)
        raise JobOpsError("DOCUMENT_OUTPUT_BUILD_FAILED", "The generated document could not be committed atomically.") from exc


def _add_numbering(document) -> int:
    _, _, _, _, OxmlElement, qn, _, _, _ = _modules()
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•")
    level.append(text)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _bullet(document, text: str, num_id: int) -> None:
    _, _, _, _, OxmlElement, qn, _, Pt, _ = _modules()
    paragraph = document.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(num)
    ppr.append(numpr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    _set_run(paragraph.add_run(text), size=10.5)


def _assert_claims(claims: list[dict[str, Any]]) -> list[tuple[str, str, list[dict[str, Any]]]]:
    approved = []
    for claim in claims:
        wordings = claim.get("allowed_wording") or []
        if not wordings:
            continue
        wording = str(wordings[0])
        decision = external_use_decision(claim, wording=wording)
        if decision.allowed:
            approved.append((str(claim["claim_id"]), wording, list(claim["source_refs"])))
    if not approved:
        raise JobOpsError("NO_APPROVED_CLAIMS", "Material composition requires at least one current approved claim.")
    return approved


def _assert_external_claims(value: dict[str, Any], *, use: str) -> list[tuple[str, str, list[dict[str, Any]]]]:
    from .external_claims import approved_external_claims

    return [
        (str(item["claim_id"]), str(item["allowed_wording"][0]), list(item.get("source_bindings", [])))
        for item in approved_external_claims(value, use=use)
    ]


def build_resume(path: Path, *, candidate_display_name: str, target_role: str, summary: str, claims: list[dict[str, Any]], skills: list[str], education: str, bullet_claim_ids: list[str] | None = None) -> dict[str, object]:
    Document, _, _, WD_ALIGN_PARAGRAPH, _, _, _, Pt, _ = _modules()
    approved = _assert_claims(claims)
    approved_text = {item[1] for item in approved}
    if summary not in approved_text:
        raise JobOpsError("SUMMARY_NOT_CLAIM_GATED", "Resume summary must use an exact approved claim wording.")
    unsupported_skills = [skill for skill in skills if skill not in approved_text]
    if unsupported_skills:
        raise JobOpsError("SKILL_NOT_CLAIM_GATED", "Every resume skill must use an exact approved claim wording.", unsupported=unsupported_skills)
    if education not in approved_text:
        raise JobOpsError("EDUCATION_NOT_CLAIM_GATED", "Resume education must use an exact approved claim wording.")
    approved_by_id = {item[0]: item for item in approved}
    selected_ids = bullet_claim_ids or [item[0] for item in approved]
    missing_ids = [claim_id for claim_id in selected_ids if claim_id not in approved_by_id]
    if missing_ids:
        raise JobOpsError("BULLET_CLAIM_NOT_APPROVED", "Every selected resume bullet requires an approved claim ID.", claim_ids=missing_ids)
    document = Document()
    _configure_document(document)
    title = document.add_paragraph(style="JobOps Title")
    title.add_run(candidate_display_name)
    subtitle = document.add_paragraph(style="JobOps Subtitle")
    subtitle.add_run(target_role)
    heading = document.add_paragraph(style="Heading 2")
    heading.add_run("Professional summary")
    p = document.add_paragraph()
    _set_run(p.add_run(summary), size=10.5)
    heading = document.add_paragraph(style="Heading 2")
    heading.add_run("Selected evidence")
    num_id = _add_numbering(document)
    for claim_id in selected_ids:
        _, wording, _ = approved_by_id[claim_id]
        _bullet(document, wording, num_id)
    heading = document.add_paragraph(style="Heading 2")
    heading.add_run("Verified skills")
    p = document.add_paragraph()
    _set_run(p.add_run(" | ".join(skills)), size=10.5)
    heading = document.add_paragraph(style="Heading 2")
    heading.add_run("Education")
    p = document.add_paragraph()
    _set_run(p.add_run(education), size=10.5)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(footer.add_run("Prepared from applicant-approved evidence"), size=8.5, color=GRAY)
    _save_document_atomic(document, path)
    return {"path": str(path), "claim_ids": [item[0] for item in approved], "preset": "compact_reference_guide", "named_overrides": []}


def build_cover_letter_narrative(
    *, candidate_display_name: str, company: str, target_role: str,
    why_company: str, why_role: str, claims: list[dict[str, Any]] | None = None,
    external_claim_set: dict[str, Any] | None = None,
) -> CoverLetterNarrative:
    if (claims is None) == (external_claim_set is None):
        raise JobOpsError("MATERIAL_CLAIM_SOURCE_INVALID", "Choose exactly one approved Claim source for the Cover Letter.")
    if claims is not None:
        approved = _assert_claims(claims)
    else:
        cover_approved = _assert_external_claims(external_claim_set or {}, use="cover_letter")
        narrative_approved = _assert_external_claims(
            external_claim_set or {}, use="application_narrative",
        )
        narrative_bindings = {
            (claim_id, wording): sources
            for claim_id, wording, sources in narrative_approved
        }
        approved = [
            (claim_id, wording, sources)
            for claim_id, wording, sources in cover_approved
            if (claim_id, wording) in narrative_bindings
        ]
        if not approved:
            raise JobOpsError(
                "APPLICATION_NARRATIVE_CLAIM_NOT_APPROVED",
                "The generated Cover Letter narrative requires the same exact Claim wording to be approved for both Cover Letter and application narrative use.",
            )
    role_claim = next((item for item in approved if item[1] == why_role), None)
    if role_claim is None:
        raise JobOpsError("WHY_ROLE_NOT_CLAIM_GATED", "Why Role must use an exact approved claim wording.")
    if "https://" not in why_company:
        raise JobOpsError("WHY_COMPANY_SOURCE_MISSING", "Why Company must include a dated HTTPS source citation.")
    evidence_claims = approved[:2]
    wording = "; ".join(item[1] for item in evidence_claims)
    used_claim_ids = tuple(dict.fromkeys([
        role_claim[0],
        *(item[0] for item in evidence_claims),
    ]))
    paragraphs = (
        "Dear Hiring Team,",
        why_company,
        why_role,
        f"Relevant verified evidence: {wording}.",
        "Thank you for considering this evidence-based application.",
        f"Sincerely,\n{candidate_display_name}",
    )
    text = validate_application_narrative_text("\n\n".join(paragraphs))
    return CoverLetterNarrative(
        candidate_display_name=candidate_display_name,
        company=company,
        target_role=target_role,
        paragraphs=paragraphs,
        text=text,
        claim_ids=used_claim_ids,
    )


def build_cover_letter(
    path: Path, *, candidate_display_name: str, company: str, target_role: str,
    why_company: str, why_role: str, claims: list[dict[str, Any]] | None = None,
    external_claim_set: dict[str, Any] | None = None,
    narrative: CoverLetterNarrative | None = None,
) -> dict[str, object]:
    Document, _, _, WD_ALIGN_PARAGRAPH, _, _, _, Pt, _ = _modules()
    expected = build_cover_letter_narrative(
        candidate_display_name=candidate_display_name,
        company=company,
        target_role=target_role,
        why_company=why_company,
        why_role=why_role,
        claims=claims,
        external_claim_set=external_claim_set,
    )
    canonical = narrative or expected
    if canonical != expected:
        raise JobOpsError(
            "APPLICATION_NARRATIVE_BINDING_INVALID",
            "The Cover Letter document and application narrative do not share one canonical source.",
        )
    document = Document()
    _configure_document(document)
    title = document.add_paragraph(style="JobOps Letter Title")
    title.add_run("APPLICATION LETTER")
    metadata = document.add_paragraph(style="JobOps Letter Subtitle")
    metadata.add_run(f"{target_role} | {company}")
    for text in canonical.paragraphs[:3]:
        paragraph = document.add_paragraph()
        _set_run(paragraph.add_run(text), size=11)
    evidence = document.add_paragraph()
    evidence.paragraph_format.space_after = Pt(10)
    _set_run(evidence.add_run(canonical.paragraphs[3]), size=11)
    closing = document.add_paragraph()
    _set_run(closing.add_run(canonical.paragraphs[4]), size=11)
    signoff = document.add_paragraph()
    signoff.paragraph_format.space_before = Pt(12)
    _set_run(signoff.add_run(canonical.paragraphs[5]), size=11)
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(footer.add_run("Prepared from applicant-approved evidence"), size=8.5, color=GRAY)
    _save_document_atomic(document, path)
    return {"path": str(path), "claim_ids": list(canonical.claim_ids), "preset": "compact_reference_guide", "named_overrides": []}


def export_docx_to_pdf(docx_path: Path, pdf_path: Path, powershell_script: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_pdf = pdf_path.with_name(f".{pdf_path.stem}.jobflow-{uuid.uuid4().hex}.tmp.pdf")
    command = [
        "powershell.exe", "-NoLogo", "-NoProfile", "-NonInteractive",
        "-ExecutionPolicy", "Bypass", "-File", str(powershell_script),
        "-InputPath", str(docx_path), "-OutputPath", str(temporary_pdf),
    ]
    try:
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        if completed.returncode != 0 or not temporary_pdf.is_file() or temporary_pdf.stat().st_size == 0:
            raise JobOpsError("DOCX_PDF_EXPORT_FAILED", "Microsoft Word PDF export failed.", returncode=completed.returncode, stderr=completed.stderr[-1000:])
        os.replace(temporary_pdf, pdf_path)
    except JobOpsError:
        temporary_pdf.unlink(missing_ok=True)
        raise
    except (OSError, subprocess.SubprocessError) as exc:
        temporary_pdf.unlink(missing_ok=True)
        raise JobOpsError("DOCX_PDF_EXPORT_FAILED", "Microsoft Word PDF export could not be committed safely.") from exc


def render_pdf_to_pngs(pdf_path: Path, output_dir: Path, pdftoppm: str) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / f".{pdf_path.stem}.jobflow-{uuid.uuid4().hex}"
    try:
        completed = subprocess.run(
            [pdftoppm, "-png", "-r", "150", str(pdf_path), str(prefix)],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
    except (OSError, subprocess.SubprocessError) as exc:
        for page in output_dir.glob(f"{prefix.name}-*.png"):
            page.unlink(missing_ok=True)
        raise JobOpsError("PDF_RENDER_FAILED", "Poppler could not start a bounded PDF render.") from exc
    pages = sorted(
        output_dir.glob(f"{prefix.name}-*.png"),
        key=lambda page: int(page.stem.rsplit("-", 1)[-1]) if page.stem.rsplit("-", 1)[-1].isdigit() else 10**9,
    )
    if completed.returncode != 0 or not pages or any(not page.is_file() or page.stat().st_size < 1 for page in pages):
        for page in pages:
            page.unlink(missing_ok=True)
        raise JobOpsError("PDF_RENDER_FAILED", "Poppler did not produce fresh page images.", returncode=completed.returncode, stderr=completed.stderr[-1000:])
    return pages
