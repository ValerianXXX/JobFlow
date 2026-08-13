from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("complex-master-resume.docx")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    for style_name, size, color in (("Title", 22, "17365D"), ("Heading 1", 12, "1F4E79"), ("Heading 2", 10.5, "365F91")):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SYNTHETIC MASTER RESUME FIXTURE")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Complex-layout format-preservation fixture | Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("{{CANDIDATE_NAME}}")
    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.add_run("{{TARGET_ROLE}}").bold = True
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(contact, "portfolio.example", "https://portfolio.example.test")

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Light Shading Accent 1"
    summary_table.columns[0].width = Inches(1.35)
    summary_table.columns[1].width = Inches(5.8)
    summary_table.cell(0, 0).text = "PROFILE"
    summary_table.cell(0, 1).text = "{{SUMMARY}}"
    for cell in summary_table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("EXPERIENCE", style="Heading 1")
    meta = doc.add_table(rows=1, cols=2)
    meta.autofit = False
    meta.columns[0].width = Inches(5.5)
    meta.columns[1].width = Inches(1.65)
    meta.cell(0, 0).text = "Synthetic Analytics Studio | Senior Analyst"
    meta.cell(0, 1).text = "2023—Present"
    meta.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    bullet = doc.add_paragraph(style="List Bullet")
    bullet.add_run("{{EXPERIENCE_BULLET}}")

    doc.add_paragraph("PROJECTS", style="Heading 1")
    project = doc.add_paragraph()
    project.paragraph_format.left_indent = Inches(0.18)
    project.paragraph_format.first_line_indent = Inches(-0.18)
    project.add_run("Selected project — ").bold = True
    project.add_run("{{PROJECT}}")

    two_col = doc.add_table(rows=2, cols=2)
    two_col.style = "Light Grid Accent 1"
    two_col.cell(0, 0).text = "SKILLS"
    two_col.cell(0, 1).text = "EDUCATION"
    two_col.cell(1, 0).text = "{{SKILLS}}"
    two_col.cell(1, 1).text = "{{EDUCATION}}"

    doc.core_properties.author = "Synthetic Fixture Builder"
    doc.core_properties.last_modified_by = "Synthetic Fixture Builder"
    doc.core_properties.comments = "Synthetic fixture only"
    doc.save(OUT)


if __name__ == "__main__":
    main()
