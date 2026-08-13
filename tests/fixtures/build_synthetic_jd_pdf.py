from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "synthetic-forward-jd.docx"


def main() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    document.add_heading("Synthetic Data Analyst", 0)
    document.add_paragraph("Example Analytics Lab | Remote")
    document.add_heading("Responsibilities", level=1)
    document.add_paragraph("Analyze synthetic datasets using Python and SQL.", style="List Bullet")
    document.add_paragraph("Document reproducible checks and communicate findings.", style="List Bullet")
    document.add_heading("Requirements", level=1)
    document.add_paragraph("At least 1 year of synthetic analytics experience.", style="List Bullet")
    document.add_paragraph("Python and SQL are required.", style="List Bullet")
    document.add_paragraph("This is a local test fixture. It is not a real job listing.")
    document.core_properties.author = "JobOps"
    document.core_properties.last_modified_by = "JobOps"
    document.save(OUTPUT)


if __name__ == "__main__":
    main()
