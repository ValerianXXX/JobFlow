from __future__ import annotations

import json
import shutil
import sys
import tempfile
import threading
import time
import unittest
from pathlib import Path
from unittest import mock

from docx import Document

from _support import PROJECT, project_temp
from jobops.db import JobOpsDB
from jobops.browser_assist import COMPANION_EXTENSION_ORIGIN
from jobops.continuous_intake import ContinuousIntakeDescriptorStore, run_continuous_intake_tick
from jobops.document_builder import inspect_docx_text_blocks, template_fingerprint
from jobops.errors import JobOpsError
from jobops.ai_runtime import LocalSubprocessAIEngine
from jobops.external_claims import build_external_claim_set, claim_review_hash
from jobops.adapters import audit_real_external_actions
from jobops.orchestrator import JobOpsOrchestrator
from jobops.onboarding_center import OnboardingCenterService
from jobops.private_onboarding import PrivateOnboarding
from jobops.queue_manager import QueueManager
from jobops.resume_tailoring import build_resume_tailoring_manifest, build_tailoring_proposal
from jobops.secure_store import WindowsDPAPIStore
from jobops.util import canonical_json, iso_utc, parse_iso, sha256_bytes, sha256_file, stable_id


SYNTHETIC_PHONE = "+1 555 010 0" + "200"
SYNTHETIC_STREET_ADDRESS = "100 Example " + "Avenue"


def manual_navigation_evidence(service, token: str, completed: dict) -> dict:
    challenge = completed["manual_navigation"]["challenge"]
    lease = service.browser_assist._leases[token]
    return {
        "trusted_user_event": True,
        "event_hash": service.browser_assist._manual_event_hash(lease),
        "prior_page_content_hash": completed["manual_navigation"]["prior_page_content_hash"],
        "control_semantics_hash": completed["manual_navigation"]["control_semantics_hash"],
        "manual_navigation_challenge_id": challenge["challenge_id"],
        "manual_navigation_nonce": challenge["nonce"],
        "manual_navigation_challenge_hash": challenge["challenge_hash"],
        "manual_navigation_assist_id": challenge["assist_id"],
        "manual_navigation_application_id": challenge["application_id"],
        "manual_navigation_tab_id": challenge["tab_id"],
        "manual_navigation_document_id": challenge["document_instance_id"],
        "manual_navigation_stage": challenge["stage"],
        "manual_navigation_client_ref": challenge["client_ref"],
        "manual_navigation_default_prevented": False,
    }


class RealProfileOfflineApplicationTests(unittest.TestCase):
    def build(self, temp: Path):
        database = JobOpsDB(temp / "jobops.db")
        database.initialize()
        private_temp = Path(tempfile.mkdtemp(prefix="jobflow-private-real-profile-test-"))
        self.addCleanup(shutil.rmtree, private_temp, True)
        onboarding = PrivateOnboarding(
            database,
            WindowsDPAPIStore(
                PROJECT / ".agents" / "skills" / "job-application-operator" / "scripts" / "secure-store.ps1",
                local_app_data=private_temp,
            ),
        )
        return database, onboarding, JobOpsOrchestrator(PROJECT, database, onboarding)

    def seed_completed_context(
        self,
        onboarding: PrivateOnboarding,
        *,
        master_kind: str = "master_resume_docx",
    ) -> dict[str, str]:
        fixtures = PROJECT / "tests" / "fixtures"
        portfolio = onboarding.import_file("onboarding_source_document", fixtures / "synthetic-forward-jd.pdf", synthetic=False)
        profile = json.loads((fixtures / "synthetic-forward-profile.json").read_text(encoding="utf-8"))
        profile.update({
            "first_name": "Synthetic",
            "last_name": "Candidate",
            "email": "synthetic-candidate@example.test",
            "phone": SYNTHETIC_PHONE,
            "phone_type": "Mobile",
            "address": SYNTHETIC_STREET_ADDRESS,
            "city": "New York",
            "state": "NY",
            "postal_code": "10001",
            "country": "United States",
            "github_url": "https://github.com/synthetic-candidate",
            "portfolio_url": "https://portfolio.example.test/synthetic-candidate",
            "portfolio_file_ref": portfolio["secure_ref"],
            "portfolio_file_sha256": portfolio["content_sha256"],
            "portfolio_file_display_name": "portfolio-fixture.pdf",
        })
        profile_record = onboarding.import_bytes("candidate_profile", canonical_json(profile), synthetic=False)
        profile_ref = str(profile_record["secure_ref"])
        profile["profile_ref"] = profile_ref
        onboarding.rotate(profile_ref, canonical_json(profile))

        answers = {
            "schema_version": 2, "status": "ONBOARDING_COMPLETE", "locale": "en", "answers": {},
            "completion": {"total": 25, "resolved": 25, "remaining": 0, "remaining_fields": [], "percent": 100},
            "updated_at": iso_utc(),
        }
        answer_record = onboarding.import_bytes("answer_bank", canonical_json(answers), synthetic=False)
        approval_record = onboarding.import_bytes("claim_approvals", canonical_json({"status": "ONBOARDING_COMPLETE"}), synthetic=False)
        master_source = Path(tempfile.mkdtemp(prefix="jobflow-real-master-source-"))
        self.addCleanup(shutil.rmtree, master_source, True)
        master_path = master_source / "reviewed-master.docx"
        document = Document()
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph("Operations analyst focused on evidence-based application review.")
        document.add_heading("Projects", level=1)
        document.add_paragraph("Analyzed synthetic datasets with reproducible methods for application review.")
        document.add_heading("Education", level=1)
        document.add_paragraph("Synthetic University, Bachelor of Science.")
        document.save(master_path)
        master_record = onboarding.import_file(master_kind, master_path, synthetic=False)
        master_ref = str(master_record["secure_ref"])
        master_hash = str(master_record["content_sha256"])
        fingerprint_hash = sha256_bytes(canonical_json(template_fingerprint(master_path).as_dict()))
        state_value: dict[str, object] = {"status": "ONBOARDING_COMPLETE"}
        if master_kind == "onboarding_source_document":
            state_value["master_resume"] = {
                "secure_ref": master_ref,
                "sha256": master_hash,
                "extension": ".docx",
                "source_id": "SRC-MASTER-01",
                "editable_docx": True,
                "template_fingerprint": fingerprint_hash,
                "template_slots": [],
                "designated_at": iso_utc(),
            }
        state_record = onboarding.import_bytes(
            "onboarding_center_state", canonical_json(state_value), synthetic=False,
        )
        state_ref = str(state_record["secure_ref"])

        block = next(item for item in inspect_docx_text_blocks(master_path) if item["text"].startswith("Analyzed synthetic"))
        claim = {
            "claim_id": "CLM-REAL-OFFLINE-01", "category": "project", "claim_kind": "achievement",
            "statement": "Analyzes synthetic datasets with reproducible methods for reviewed applications.",
            "decision": "CONFIRMED", "deleted": False, "source_id": "SRC-MASTER-01",
            "provenance": {"line_start": block["line_number"], "line_end": block["line_number"]},
            "source_bindings": [{"kind": "MASTER_RESUME", "secure_ref": master_ref, "content_sha256": master_hash}],
        }
        master = {
            "secure_ref": master_ref, "sha256": master_hash, "editable_docx": True,
            "template_fingerprint": fingerprint_hash, "source_id": "SRC-MASTER-01",
        }
        external = build_external_claim_set(
            onboarding_state_ref=state_ref, profile_ref=profile_ref, master_resume=master,
            claims=[claim], allowed_uses=["resume", "cover_letter", "application_narrative"],
            expected_review_hash=claim_review_hash([claim], master_hash),
        )
        external_record = onboarding.import_bytes("external_claim_set", canonical_json(external), synthetic=False)
        proposal = build_tailoring_proposal(
            onboarding_state_ref=state_ref, master_resume=master,
            blocks=inspect_docx_text_blocks(master_path), claims=[claim],
        )
        candidate = proposal["candidates"][0]
        manifest = build_resume_tailoring_manifest(
            onboarding_state_ref=state_ref, master_resume=master, proposal=proposal,
            selections=[{"block_ref": candidate["block_ref"], "category": "project"}],
            expected_proposal_hash=proposal["proposal_hash"], user_confirmed=True,
        )
        manifest_record = onboarding.import_bytes("resume_tailoring_manifest", canonical_json(manifest), synthetic=False)
        completion = {
            "schema_version": 1, "status": "ONBOARDING_COMPLETE", "profile_ref": profile_ref,
            "answer_bank_ref": answer_record["secure_ref"], "claim_approvals_ref": approval_record["secure_ref"],
            "counts": {"answers_resolved": 25, "answers_total": 25, "claims_reviewed": 1, "claims_total": 1, "conflicts_resolved": 0, "conflicts_total": 0},
            "sources": {"resume_or_material": 1, "ai": 0, "direct_answers": 25},
            "locale": "en", "completed_at": iso_utc(), "real_external_actions": 0,
            "knowledge_write_operations": 0,
        }
        onboarding.import_bytes("onboarding_completion_packet", canonical_json(completion), synthetic=False)
        return {
            "profile_ref": profile_ref, "answer_bank_ref": str(answer_record["secure_ref"]),
            "master_resume_ref": master_ref, "external_claim_set_ref": str(external_record["secure_ref"]),
            "tailoring_manifest_ref": str(manifest_record["secure_ref"]),
            "tailoring_manifest_hash": str(manifest["content_hash"]),
        }

    def test_ui_designated_docx_master_is_accepted_only_through_its_approved_state_binding(self) -> None:
        with project_temp() as temp:
            _, onboarding, orchestrator = self.build(temp)
            references = self.seed_completed_context(
                onboarding,
                master_kind="onboarding_source_document",
            )
            resolved = orchestrator.current_real_application_references()
            self.assertEqual(resolved["master_resume_ref"], references["master_resume_ref"])
            self.assertEqual(resolved["external_claim_set_ref"], references["external_claim_set_ref"])

            claim_set = json.loads(onboarding.read_bytes(resolved["external_claim_set_ref"]))
            replacement_ref = onboarding.import_bytes(
                "onboarding_source_document", b"unapproved replacement", synthetic=False,
            )["secure_ref"]
            onboarding.rotate(
                str(claim_set["onboarding_state_ref"]),
                canonical_json({
                    "status": "ONBOARDING_COMPLETE",
                    "master_resume": {
                        "secure_ref": replacement_ref,
                        "sha256": "sha256:" + "0" * 64,
                        "extension": ".docx",
                        "editable_docx": True,
                    },
                }),
            )
            with self.assertRaises(JobOpsError) as rejected:
                orchestrator.current_real_application_references()
            self.assertEqual(rejected.exception.code, "APPLICATION_MASTER_SOURCE_BINDING_INVALID")

    def test_ai_operator_material_and_form_turns_are_bound_to_one_real_preparation(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            engine = LocalSubprocessAIEngine([
                sys.executable,
                str(PROJECT / "tests" / "fixtures" / "fake_jobops_ai.py"),
            ])
            service = OnboardingCenterService(
                PROJECT, database, onboarding, ai_engine=engine,
            )
            self.addCleanup(service.close)
            operator_task_id = "AIT-0123456789AB"
            fixtures = PROJECT / "tests" / "fixtures"
            result = service.prepare_offline_application_bundle(
                metadata={
                    "official_url": "https://example.com/careers/synthetic-data-analyst",
                    "application_url": "https://example.com/careers/apply/synthetic-data-analyst",
                    "guest_available": True,
                    "research_title": "Synthetic company role page",
                    "evidence_excerpt": "Synthetic Data Analyst",
                    "operator_task_id": operator_task_id,
                },
                files={
                    "jd": (".txt", (fixtures / "synthetic-forward-jd.txt").read_bytes()),
                    "official": (".html", b"<h1>Synthetic Data Analyst</h1><a href='/careers/apply/synthetic-data-analyst'>Apply</a>"),
                    "form": (".html", (fixtures / "synthetic-material-form.html").read_bytes()),
                },
            )
            application_id = str(result["application_id"])
            packet = service.review_packet(application_id)["packet"]
            operator = packet["ai_operator"]

            self.assertEqual(operator["operator_task_id"], operator_task_id)
            self.assertEqual(len(operator["turns"]), 2)
            self.assertEqual(
                [turn["decision_point"] for turn in operator["turns"]],
                ["JOB_AND_MATERIAL_DECISION", "CURRENT_FORM_SEMANTIC_REVIEW"],
            )
            self.assertTrue(all(turn["application_id"] == application_id for turn in operator["turns"]))
            self.assertTrue(all(turn["task_id"] == operator_task_id for turn in operator["turns"]))
            self.assertTrue(all(turn["final_submit"] == "USER_ONLY" for turn in operator["turns"]))
            self.assertTrue(all(turn["automatic_retry"] is False for turn in operator["turns"]))
            self.assertTrue(all(turn["private_values_exposed"] == 0 for turn in operator["turns"]))

            activity = service.bootstrap()["ai_operator"]["activity"]
            relevant = [item for item in activity["recent_turns"] if item["application_id"] == application_id]
            self.assertEqual(len(relevant), 2)
            self.assertTrue(all(item["status"] == "HOST_PIPELINE_VERIFIED" for item in relevant))
            serialized = json.dumps({"operator": operator, "activity": relevant})
            self.assertNotIn("synthetic-candidate@example.test", serialized)
            self.assertNotIn(SYNTHETIC_PHONE, serialized)
            self.assertNotIn(SYNTHETIC_STREET_ADDRESS, serialized)

            decision = service.decide_review_packet({
                "application_id": application_id,
                "decision": "APPROVE",
                "expected_packet_hash": packet["content_hash"],
                "user_confirmed": True,
            })
            self.assertEqual(decision["status"], "APPROVED")
            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            self.assertEqual(bundle["operator_task_id"], operator_task_id)
            _token, prepared = self.prepared_browser_assist(service, application_id)
            live_turn = prepared["ai_operator"]
            self.assertEqual(live_turn["task_id"], operator_task_id)
            self.assertEqual(live_turn["application_id"], application_id)
            self.assertEqual(live_turn["selected_tool"], "jobflow.inspect_application_form")
            self.assertEqual(live_turn["status"], "HOST_PIPELINE_VERIFIED")
            self.assertEqual(live_turn["final_submit"], "USER_ONLY")
            self.assertFalse(live_turn["automatic_retry"])
            self.assertEqual(live_turn["private_values_exposed"], 0)
            self.assertEqual(live_turn["real_external_actions"], 0)
            with database.connect() as connection:
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM events WHERE application_id=? AND event_type='AI_OPERATOR_TURN'",
                    (application_id,),
                ).fetchone()[0], 3)
            safe_live = json.dumps(live_turn)
            self.assertNotIn("synthetic-candidate@example.test", safe_live)
            self.assertNotIn(SYNTHETIC_PHONE, safe_live)
            self.assertNotIn(SYNTHETIC_STREET_ADDRESS, safe_live)
            stopped = service.browser_assist.stop(user_confirmed=True)
            self.assertEqual(stopped["status"], "BROWSER_ASSIST_STOPPED")

    def approved_company_application(
        self, database: JobOpsDB, onboarding: PrivateOnboarding,
    ) -> tuple[OnboardingCenterService, str, dict]:
        fixtures = PROJECT / "tests" / "fixtures"
        service = OnboardingCenterService(PROJECT, database, onboarding)
        self.addCleanup(service.close)
        result = service.prepare_offline_application_bundle(
            metadata={
                "official_url": "https://example.com/careers/synthetic-data-analyst",
                "application_url": "https://example.com/careers/apply/synthetic-data-analyst",
                "guest_available": True,
                "research_title": "Synthetic company role page",
                "evidence_excerpt": "Synthetic Data Analyst",
            },
            files={
                "jd": (".txt", (fixtures / "synthetic-forward-jd.txt").read_bytes()),
                "official": (".html", b"<html><body><h1>Synthetic Data Analyst</h1><a href='https://example.com/careers/apply/synthetic-data-analyst'>Apply</a></body></html>"),
                "form": (".html", (fixtures / "synthetic-material-form.html").read_bytes()),
            },
        )
        application_id = str(result["application_id"])
        packet = service.review_packet(application_id)["packet"]
        decision = service.decide_review_packet({
            "application_id": application_id,
            "decision": "APPROVE",
            "expected_packet_hash": packet["content_hash"],
            "user_confirmed": True,
        })
        self.assertEqual(decision["status"], "APPROVED")
        return service, application_id, packet

    def test_resume_contact_fields_are_ready_without_reasking_in_single_review(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            form = b"""<!doctype html><html><body><form action='/careers/apply/contact-ready'>
                <label for='first-name'>First name</label><input id='first-name' name='first_name' type='text' required>
                <label for='last-name'>Last name</label><input id='last-name' name='last_name' type='text' required>
                <label for='email'>Email</label><input id='email' name='email' type='email' required>
                <label for='phone'>Phone</label><input id='phone' name='phone' type='tel' required>
                <label for='phone-type'>Phone Type</label><select id='phone-type' name='phone_type' required>
                  <option value='Mobile'>Mobile</option><option value='Home'>Home</option>
                </select>
                <label for='address'>Mailing Address</label><input id='address' name='mailing_address' type='text' required>
                <label for='city'>City</label><input id='city' name='city' type='text' required>
                <label for='state'>State</label><select id='state' name='state' required>
                  <option value='NY'>New York</option><option value='NJ'>New Jersey</option>
                </select>
                <label for='postal-code'>Zip Code</label><input id='postal-code' name='postal_code' type='text' required>
                <label for='country'>Country</label><select id='country' name='country' required>
                  <option value='United States'>United States</option><option value='Canada'>Canada</option>
                </select>
                <button type='submit'>Submit application</button>
            </form></body></html>"""
            prepared = service.prepare_offline_application_bundle(
                metadata={
                    "official_url": "https://example.com/careers/contact-ready",
                    "application_url": "https://example.com/careers/apply/contact-ready",
                    "guest_available": True,
                    "research_title": "Synthetic contact-ready role",
                    "evidence_excerpt": "Synthetic Data Analyst",
                },
                files={
                    "jd": (".txt", (PROJECT / "tests" / "fixtures" / "synthetic-forward-jd.txt").read_bytes()),
                    "official": (".html", b"<h1>Synthetic Data Analyst</h1><a href='/careers/apply/contact-ready'>Apply</a>"),
                    "form": (".html", form),
                },
            )
            reviewed = service.review_packet(str(prepared["application_id"]))
            self.assertEqual(reviewed["field_resolution"]["unresolved_count"], 0)
            questions = {
                str(item["answer_key"]): item
                for item in reviewed["packet"]["form_questions"]
                if item.get("answer_key") in {
                    "first_name", "last_name", "email", "phone", "phone_type",
                    "address", "city", "state", "postal_code", "country",
                }
            }
            self.assertEqual(set(questions), {
                "first_name", "last_name", "email", "phone", "phone_type",
                "address", "city", "state", "postal_code", "country",
            })
            self.assertTrue(all(item["status"] == "READY" for item in questions.values()))
            self.assertTrue(all(item["redacted_summary"] == "PRIVATE_VALUE_PRESENT" for item in questions.values()))
            serialized = json.dumps(reviewed)
            self.assertNotIn("synthetic-candidate@example.test", serialized)
            self.assertNotIn(SYNTHETIC_PHONE, serialized)
            self.assertNotIn(SYNTHETIC_STREET_ADDRESS, serialized)

    def approved_workday_v2_application(
        self, database: JobOpsDB, onboarding: PrivateOnboarding,
        *, initial_form: str = "synthetic-v2-workday-step-1.html",
    ) -> tuple[OnboardingCenterService, str, dict]:
        fixtures = PROJECT / "tests" / "fixtures"
        service = OnboardingCenterService(PROJECT, database, onboarding)
        self.addCleanup(service.close)
        official = (
            "<html><body><h1>Synthetic Data Analyst</h1>"
            "<a href='https://example.wd5.myworkdayjobs.com/en-US/Careers/job/123'>Apply on Workday</a>"
            "</body></html>"
        ).encode("utf-8")
        result = service.prepare_offline_application_bundle(
            metadata={
                "official_url": "https://example.com/careers/synthetic-data-analyst",
                "application_url": "https://example.wd5.myworkdayjobs.com/en-US/Careers/job/123",
                "guest_available": True,
                "research_title": "Synthetic company role page",
                "evidence_excerpt": "Synthetic Data Analyst",
            },
            files={
                "jd": (".txt", (fixtures / "synthetic-forward-jd.txt").read_bytes()),
                "official": (".html", official),
                "form": (".html", (fixtures / initial_form).read_bytes()),
            },
        )
        application_id = str(result["application_id"])
        packet = service.review_packet(application_id)["packet"]
        decision = service.decide_review_packet({
            "application_id": application_id,
            "decision": "APPROVE",
            "expected_packet_hash": packet["content_hash"],
            "user_confirmed": True,
        })
        self.assertEqual(decision["status"], "APPROVED")
        return service, application_id, packet

    def prepared_browser_assist(
        self, service: OnboardingCenterService, application_id: str,
    ) -> tuple[str, dict]:
        started = service.start_browser_assist({
            "application_id": application_id, "user_confirmed": True,
        })
        resumed = service.start_browser_assist({
            "application_id": application_id, "user_confirmed": True,
        })
        self.assertTrue(resumed["resumed"])
        self.assertEqual(resumed["assist_id"], started["assist_id"])
        self.assertEqual(resumed["assist_token"], started["assist_token"])
        token = str(started["assist_token"])
        paired = service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
        self.assertEqual(paired["status"], "BROWSER_COMPANION_PAIRED")
        self.assertEqual(paired["capture_status"], "READY")
        paired_again = service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
        self.assertEqual(paired_again["assist_id"], paired["assist_id"])
        with service.database.connect() as connection:
            self.assertEqual(connection.execute(
                "SELECT COUNT(*) FROM browser_assist_events WHERE assist_id=? AND event_type='COMPANION_PAIRED'",
                (paired["assist_id"],),
            ).fetchone()[0], 1)
            self.assertEqual(connection.execute(
                "SELECT COUNT(*) FROM browser_assist_events WHERE assist_id=? AND event_type='ASSIST_STARTED'",
                (paired["assist_id"],),
            ).fetchone()[0], 1)
        bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
        field_count = len(bundle["form_snapshot"]["fields"])
        prepared = service.browser_assist.prepare(
            token,
            {
                "url": started["approved_url"],
                "sanitized_html": (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8"),
                "client_refs": [f"DOM-{index:012d}" for index in range(1, field_count + 1)],
                "blocker_signals": [],
            },
            extension_origin=COMPANION_EXTENSION_ORIGIN,
        )
        return token, prepared

    def test_guided_browser_intake_builds_one_review_packet_without_page_actions(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(
                onboarding,
                master_kind="onboarding_source_document",
            )
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            fixtures = PROJECT / "tests" / "fixtures"
            readiness = {"application_readiness": {"status": "READY_FOR_OFFLINE_APPLICATION_PREPARATION"}}
            with mock.patch.object(service, "bootstrap", return_value=readiness):
                started = service.start_guided_intake({
                    "official_url": "https://example.com/careers/synthetic-data-analyst",
                    "user_confirmed": True,
                })
            token = str(started["intake_token"])
            paired = service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(paired["capture_status"], "AWAITING_JOB_PAGE_CAPTURE")
            job_text = (fixtures / "synthetic-forward-jd.txt").read_text(encoding="utf-8")
            captured = service.capture_guided_job_page(
                token,
                {
                    "url": "https://example.com/careers/synthetic-data-analyst",
                    "document_title": "Synthetic Data Analyst",
                    "job_title": "Synthetic Data Analyst",
                    "company_name": "Example Analytics Lab",
                    "job_location": "Remote",
                    "visible_text": job_text,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(captured["status"], "AWAITING_APPLICATION_FORM_CAPTURE")
            paired_again = service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(paired_again["capture_status"], "AWAITING_APPLICATION_FORM_CAPTURE")
            with database.connect() as connection:
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM guided_intake_events WHERE intake_id=? AND event_type='PAIRED'",
                    (paired["intake_id"],),
                ).fetchone()[0], 1)

            prepared = service.start_guided_application_form_preparation(
                token,
                {
                    "url": "https://example.com/careers/apply/synthetic-data-analyst",
                    "sanitized_html": (fixtures / "synthetic-material-form.html").read_text(encoding="utf-8"),
                    "blocker_signals": [],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(prepared["status"], "PREPARING_APPLICATION")
            deadline = time.time() + 120
            while time.time() < deadline:
                prepared = service.guided_application_form_preparation_status(
                    token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                if prepared["status"] != "PREPARING_APPLICATION":
                    break
                time.sleep(0.05)
            self.assertEqual(prepared["status"], "REVIEW_PACKET_READY")
            packet = service.review_packet(str(prepared["application_id"]))
            self.assertEqual(packet["application_id"], prepared["application_id"])
            with self.assertRaises(JobOpsError) as completed_cancel:
                service.cancel_guided_intake({
                    "intake_id": paired["intake_id"],
                    "user_confirmed": True,
                })
            self.assertEqual(completed_cancel.exception.code, "GUIDED_INTAKE_CANCEL_UNAVAILABLE")
            with database.connect() as connection:
                events = [row[0] for row in connection.execute(
                    "SELECT event_type FROM guided_intake_events ORDER BY event_id"
                ).fetchall()]
                self.assertEqual(events, [
                    "STARTED", "PAIRED", "JOB_PAGE_INSPECTED", "APPLY_ROUTE_INSPECTED",
                    "FORM_INSPECTED", "REVIEW_PACKET_READY",
                ])
                with self.assertRaises(Exception):
                    connection.execute("UPDATE guided_intake_events SET event_type='FAILED' WHERE event_id=1")
                with self.assertRaises(Exception):
                    connection.execute("DELETE FROM guided_intake_events WHERE event_id=1")
            audit = audit_real_external_actions(database)
            self.assertEqual(audit["attempt_count"], 0)
            self.assertEqual(audit["real_external_actions"], 0)

    def test_guided_browser_discovery_selects_only_an_official_company_job_page(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding, master_kind="onboarding_source_document")
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            official_url = "https://careers.example.com/us/en/job/CR-102/Credit-Risk-Analyst?source=search"
            candidate_ref = stable_id("JDC", official_url)
            engine = mock.Mock()
            engine.ready = True
            engine.public_status.return_value = {
                "status": "READY", "structured_capability_status": "VERIFIED",
            }
            engine.execute_structured_task.return_value = {
                "schema_version": 1, "status": "SELECTED",
                "ranked_candidate_refs": [candidate_ref],
                "summary": "The official company role matches the requested title and location.",
            }
            service.ai_engine = engine
            with mock.patch.object(service, "bootstrap", return_value={
                "application_readiness": {"status": "READY_FOR_OFFLINE_APPLICATION_PREPARATION"},
            }):
                started = service.start_guided_intake({
                    "search_intent": "Find credit risk analyst roles in New York",
                    "user_confirmed": True,
                })
            token = str(started["intake_token"])
            self.assertEqual(started["discovery_mode"], "VISIBLE_BROWSER_SEARCH")
            paired = service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(paired["capture_status"], "AWAITING_JOB_DISCOVERY")
            self.assertIn("official company careers jobs", paired["search_query"])
            selected = service.capture_guided_search_results(
                token,
                {
                    "search_origin": "https://www.google.com/search",
                    "results": [
                        {
                            "url": official_url,
                            "title": "Credit Risk Analyst | Example Careers",
                            "snippet": "Join the Example credit risk team in New York.",
                        },
                        {
                            "url": "https://www.indeed.com/viewjob?jk=abc",
                            "title": "Credit Risk Analyst",
                            "snippet": "Aggregator copy.",
                        },
                    ],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(selected["status"], "AWAITING_JOB_PAGE_CAPTURE")
            self.assertEqual(selected["official_url"], official_url)
            self.assertEqual(selected["allowed_company_domain"], "example.com")
            self.assertEqual(service._guided_intakes[token]["company_domain"], "example.com")
            with database.connect() as connection:
                events = [row[0] for row in connection.execute(
                    "SELECT event_type FROM guided_intake_events WHERE intake_id=? ORDER BY event_id",
                    (paired["intake_id"],),
                ).fetchall()]
            self.assertEqual(events, ["STARTED", "PAIRED", "SEARCH_RESULTS_INSPECTED"])

    def test_guided_browser_discovery_ambiguity_accepts_only_a_displayed_candidate(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding, master_kind="onboarding_source_document")
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            first_url = "https://careers.alpha-example.com/jobs/risk-analyst"
            second_url = "https://jobs.beta-example.com/careers/risk-analyst"
            engine = mock.Mock()
            engine.ready = True
            engine.public_status.return_value = {
                "status": "READY", "structured_capability_status": "VERIFIED",
            }
            engine.execute_structured_task.return_value = {
                "schema_version": 1, "status": "NEEDS_USER_SELECTION",
                "ranked_candidate_refs": [],
                "summary": "Both verified company roles match the saved goal.",
            }
            service.ai_engine = engine
            with mock.patch.object(service, "bootstrap", return_value={
                "application_readiness": {"status": "READY_FOR_OFFLINE_APPLICATION_PREPARATION"},
            }):
                started = service.start_guided_intake({
                    "search_intent": "credit risk analyst",
                    "user_confirmed": True,
                })
            token = str(started["intake_token"])
            service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            ambiguous = service.capture_guided_search_results(
                token,
                {
                    "search_origin": "https://www.google.com/search",
                    "results": [
                        {"url": first_url, "title": "Risk Analyst | Alpha Careers", "snippet": "Company role."},
                        {"url": second_url, "title": "Risk Analyst | Beta Careers", "snippet": "Company role."},
                    ],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(ambiguous["status"], "SEARCH_SELECTION_REQUIRED")
            self.assertEqual(len(ambiguous["candidate_options"]), 2)
            self.assertEqual(service._guided_public_status()["candidate_options"], ambiguous["candidate_options"])
            with self.assertRaises(JobOpsError) as invented:
                service.select_guided_search_candidate({
                    "intake_id": started["intake_id"],
                    "candidate_ref": "JDC-FFFFFFFFFFFF",
                    "user_confirmed": True,
                })
            self.assertEqual(invented.exception.code, "OFFICIAL_JOB_SELECTION_INVALID")
            chosen_ref = stable_id("JDC", second_url)
            selected = service.select_guided_search_candidate({
                "intake_id": started["intake_id"],
                "candidate_ref": chosen_ref,
                "user_confirmed": True,
            })
            self.assertEqual(selected["official_url"], second_url)
            self.assertEqual(selected["status"], "AWAITING_JOB_PAGE_CAPTURE")
            repaired_pair = service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(repaired_pair["official_url"], second_url)
            self.assertEqual(repaired_pair["capture_status"], "AWAITING_JOB_PAGE_CAPTURE")

    def test_guided_form_preparation_runs_once_in_background_and_is_pollable(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            token = "a" * 54
            intake_id = "GIN-123456789ABC"
            service._guided_intakes[token] = {
                "token": token,
                "intake_id": intake_id,
                "official_url": "https://example.com/careers/synthetic-role",
                "company_domain": "example.com",
                "started_epoch": time.time(),
                "expires_epoch": time.time() + 1800,
                "expires_at": "2099-01-01T00:00:00Z",
                "paired": True,
                "status": "AWAITING_APPLICATION_FORM_CAPTURE",
                "job_page": {"visible_text": "Synthetic role details", "title": "Synthetic role"},
            }
            entered = threading.Event()
            release = threading.Event()
            calls = []

            def slow_preparation(*args, **kwargs):
                calls.append((args, kwargs))
                entered.set()
                self.assertTrue(release.wait(5))
                return {
                    "status": "REVIEW_PACKET_READY",
                    "application_id": "APP-123456789ABC",
                    "real_external_actions": 0,
                }

            with mock.patch.object(service, "capture_guided_application_form", side_effect=slow_preparation):
                started = service.start_guided_application_form_preparation(
                    token,
                    {"url": "https://example.com/apply", "sanitized_html": "<form><input></form>"},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                self.assertEqual(started["status"], "PREPARING_APPLICATION")
                self.assertTrue(entered.wait(2))
                pending = service.guided_application_form_preparation_status(
                    token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                self.assertEqual(pending["status"], "PREPARING_APPLICATION")
                repeated = service.start_guided_application_form_preparation(
                    token,
                    {"url": "https://example.com/other", "sanitized_html": "<form><input></form>"},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                self.assertEqual(repeated["status"], "PREPARING_APPLICATION")
                self.assertEqual(len(calls), 1, "polling or repeat clicks must not start a second preparation")
                release.set()
                deadline = time.time() + 5
                result = pending
                while time.time() < deadline:
                    result = service.guided_application_form_preparation_status(
                        token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                    )
                    if result["status"] != "PREPARING_APPLICATION":
                        break
                    time.sleep(0.02)
            self.assertEqual(result["status"], "REVIEW_PACKET_READY")
            self.assertEqual(result["application_id"], "APP-123456789ABC")
            self.assertEqual(result["intake_id"], intake_id)
            self.assertFalse(result["automatic_retry"])
            self.assertEqual(result["real_external_actions"], 0)
            self.assertEqual(len(calls), 1)

    def test_guided_browser_intake_can_replace_an_unrecoverable_pairing_lease(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            readiness = {"application_readiness": {"status": "READY_FOR_OFFLINE_APPLICATION_PREPARATION"}}
            payload = {
                "official_url": "https://example.com/careers/synthetic-data-analyst",
                "user_confirmed": True,
            }
            with mock.patch.object(service, "bootstrap", return_value=readiness):
                first = service.start_guided_intake(payload)
                second = service.start_guided_intake(payload)
            self.assertNotEqual(first["intake_id"], second["intake_id"])
            with self.assertRaises(JobOpsError) as stale:
                service.pair_guided_intake(
                    str(first["intake_token"]), extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(stale.exception.code, "GUIDED_INTAKE_NOT_FOUND")
            paired = service.pair_guided_intake(
                str(second["intake_token"]), extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(paired["capture_status"], "AWAITING_JOB_PAGE_CAPTURE")
            with database.connect() as connection:
                first_events = [str(row[0]) for row in connection.execute(
                    "SELECT event_type FROM guided_intake_events WHERE intake_id=? ORDER BY event_id",
                    (first["intake_id"],),
                ).fetchall()]
                second_events = [str(row[0]) for row in connection.execute(
                    "SELECT event_type FROM guided_intake_events WHERE intake_id=? ORDER BY event_id",
                    (second["intake_id"],),
                ).fetchall()]
            self.assertEqual(first_events, ["STARTED", "FAILED"])
            self.assertEqual(second_events, ["STARTED", "PAIRED"])
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 0)

    def test_guided_browser_intake_cancel_releases_wrong_url_without_external_actions(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service = OnboardingCenterService(PROJECT, database, onboarding)
            self.addCleanup(service.close)
            readiness = {"application_readiness": {"status": "READY_FOR_OFFLINE_APPLICATION_PREPARATION"}}
            with mock.patch.object(service, "bootstrap", return_value=readiness):
                started = service.start_guided_intake({
                    "official_url": "https://example.com/careers/wrong-role",
                    "user_confirmed": True,
                })
            token = str(started["intake_token"])
            service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            service.capture_guided_job_page(
                token,
                {
                    "url": "https://example.com/careers/wrong-role",
                    "job_title": "Wrong Synthetic Role",
                    "company_name": "Example Analytics Lab",
                    "job_location": "Remote",
                    "visible_text": "Wrong Synthetic Role\nThis is readable synthetic job content for cancellation testing.",
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            retained_lease = service._guided_intakes[token]
            with self.assertRaises(JobOpsError) as missing_confirmation:
                service.cancel_guided_intake({
                    "intake_id": started["intake_id"],
                    "user_confirmed": False,
                })
            self.assertEqual(missing_confirmation.exception.code, "EXPLICIT_CONFIRMATION_REQUIRED")
            cancelled = service.cancel_guided_intake({
                "intake_id": started["intake_id"],
                "user_confirmed": True,
            })
            self.assertEqual(cancelled["status"], "GUIDED_INTAKE_CANCELLED")
            self.assertFalse(cancelled["active"])
            self.assertFalse(cancelled["already_ended"])
            self.assertEqual(retained_lease["job_page"], None)
            self.assertEqual(service._guided_public_status()["status"], "IDLE")
            with self.assertRaises(JobOpsError) as stale:
                service.pair_guided_intake(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(stale.exception.code, "GUIDED_INTAKE_NOT_FOUND")
            cancelled_again = service.cancel_guided_intake({
                "intake_id": started["intake_id"],
                "user_confirmed": True,
            })
            self.assertTrue(cancelled_again["already_ended"])
            with mock.patch.object(service, "bootstrap", return_value=readiness):
                restarted = service.start_guided_intake({
                    "official_url": "https://example.com/careers/correct-role",
                    "user_confirmed": True,
                })
            self.assertEqual(restarted["status"], "GUIDED_INTAKE_PAIRING")
            self.assertNotEqual(restarted["intake_id"], started["intake_id"])
            with database.connect() as connection:
                events = [str(row[0]) for row in connection.execute(
                    "SELECT event_type FROM guided_intake_events WHERE intake_id=? ORDER BY event_id",
                    (started["intake_id"],),
                ).fetchall()]
            self.assertEqual(events, [
                "STARTED", "PAIRED", "JOB_PAGE_INSPECTED", "APPLY_ROUTE_INSPECTED", "FAILED",
            ])
            actions = audit_real_external_actions(database)
            self.assertEqual(actions["attempt_count"], 0)
            self.assertEqual(actions["real_external_actions"], 0)

    def test_completed_user_context_generates_encrypted_review_materials_without_external_actions(self) -> None:
        with project_temp() as temp:
            database, onboarding, orchestrator = self.build(temp)
            refs = self.seed_completed_context(onboarding)
            fixtures = PROJECT / "tests" / "fixtures"
            route = json.loads((fixtures / "synthetic-forward-route.json").read_text(encoding="utf-8"))
            route["research"] = {
                "title": "Example Analytics Lab update", "url": "https://example.com/news/synthetic-update",
                "source_type": "official_company", "published_at": "2026-08-12T00:00:00Z",
                "accessed_at": iso_utc(), "official": True,
                "evidence_excerpt": "Example Analytics Lab uses documented checks for synthetic dataset analysis.",
            }
            route_path = temp / "saved-route.json"
            route_path.write_text(json.dumps(route), encoding="utf-8")
            master_before = onboarding.read_bytes(refs["master_resume_ref"])

            try:
                result = orchestrator.run_to_awaiting(
                    fixtures / "synthetic-forward-jd.txt",
                    profile_ref=None, master_resume_ref=None, answer_bank_ref=None,
                    route_fixture=route_path, form_fixture=fixtures / "synthetic-material-form.html",
                    research_fixture=fixtures / "synthetic-research.html", synthetic=False,
                )
            except JobOpsError as exc:
                details = exc.as_dict()
                if isinstance(details.get("stderr"), str):
                    details["stderr"] = details["stderr"][-800:]
                if isinstance(details.get("stdout"), str):
                    details["stdout"] = details["stdout"][-800:]
                self.fail(json.dumps(details, ensure_ascii=False))

            self.assertEqual(result["status"], "AWAITING_APPROVAL")
            self.assertEqual(result["document_qa"]["status"], "PASS")
            self.assertEqual(result["cover_letter_qa"]["status"], "PASS")
            self.assertEqual(result["real_external_actions"], 0)
            self.assertEqual(onboarding.read_bytes(refs["master_resume_ref"]), master_before)
            packet_meta = onboarding.reference_metadata(str(result["review_packet_ref"]))
            self.assertFalse(packet_meta["synthetic"])
            packet = json.loads(onboarding.read_bytes(str(result["review_packet_ref"])).decode("utf-8"))
            self.assertEqual(packet["master_resume_diff"]["manifest_content_hash"], refs["tailoring_manifest_hash"])
            self.assertTrue(packet["material_plan"]["all_uploads_and_submission_blocked"])
            self.assertFalse(packet["execution_plan"]["live_transport_registered"])
            with database.connect() as connection:
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 0)
                generated = connection.execute(
                    "SELECT COUNT(*) FROM private_refs WHERE kind LIKE 'generated_%' AND synthetic=0 AND status='ACTIVE'"
                ).fetchone()[0]
            self.assertGreaterEqual(generated, 4)
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])

    def test_local_ui_bundle_builds_route_and_review_packet_without_retaining_input_files(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            fixtures = PROJECT / "tests" / "fixtures"
            service = OnboardingCenterService(PROJECT, database, onboarding)
            result = service.prepare_offline_application_bundle(
                metadata={
                    "official_url": "https://example.com/careers/synthetic-data-analyst",
                    "application_url": "https://boards.greenhouse.io/example/jobs/987654",
                    "guest_available": True,
                    "research_title": "Synthetic company role page",
                    "evidence_excerpt": "Synthetic Data Analyst",
                },
                files={
                    "jd": (".txt", (fixtures / "synthetic-forward-jd.txt").read_bytes()),
                    "official": (".html", (fixtures / "synthetic-greenhouse-careers.html").read_bytes()),
                    "form": (".html", (fixtures / "synthetic-greenhouse-form.html").read_bytes()),
                },
            )
            self.assertEqual(result["status"], "AWAITING_APPROVAL")
            self.assertFalse(result["deferred_evidence_retained"])
            self.assertEqual(result["real_external_actions"], 0)
            self.assertEqual(result["network_actions"], 0)
            displayed = service.review_packet(str(result["application_id"]))
            self.assertEqual(displayed["packet"]["source_route"]["provider"], "greenhouse")
            with database.connect() as connection:
                kinds = {str(row[0]) for row in connection.execute("SELECT DISTINCT kind FROM private_refs")}
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 0)
            self.assertFalse(any(kind.startswith("application_input_") for kind in kinds))
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])

    def test_ui_deferred_bundle_is_encrypted_then_automatically_consumed_after_review(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            database.set_pending_limit(1)
            fixtures = PROJECT / "tests" / "fixtures"
            service = OnboardingCenterService(PROJECT, database, onboarding)
            metadata = {
                "official_url": "https://example.com/careers/synthetic-data-analyst",
                "application_url": "https://boards.greenhouse.io/example/jobs/987654",
                "guest_available": True,
                "research_title": "Synthetic company role page",
                "evidence_excerpt": "Synthetic Data Analyst",
            }

            def local_bundle(jd_name: str) -> dict[str, tuple[str, bytes]]:
                return {
                    "jd": (".txt", (fixtures / jd_name).read_bytes()),
                    "official": (".html", (fixtures / "synthetic-greenhouse-careers.html").read_bytes()),
                    "form": (".html", (fixtures / "synthetic-material-form.html").read_bytes()),
                }

            first = service.prepare_offline_application_bundle(
                metadata=metadata, files=local_bundle("synthetic-forward-jd.txt"),
            )
            second = service.prepare_offline_application_bundle(
                metadata=metadata, files=local_bundle("synthetic-forward-jd-two.txt"),
            )
            self.assertEqual(first["status"], "AWAITING_APPROVAL")
            self.assertEqual(second["status"], "DEFERRED")
            self.assertTrue(second["deferred_evidence_retained"])
            self.assertEqual(len(list((temp / "continuous-intake").glob("*.json"))), 1)
            with database.connect() as connection:
                active_before = connection.execute(
                    "SELECT COUNT(*) FROM private_refs WHERE kind='continuous_evidence_bundle' AND status='ACTIVE'"
                ).fetchone()[0]
            self.assertEqual(active_before, 1)

            application_id = str(first["application_id"])
            packet = service.review_packet(application_id)["packet"]
            decision = service.decide_review_packet({
                "application_id": application_id, "decision": "APPROVE",
                "expected_packet_hash": packet["content_hash"], "user_confirmed": True,
            })
            continued = decision["continued_intake"]
            self.assertEqual(continued["status"], "LOCAL_CONTINUATION_PROCESSED")
            self.assertEqual(continued["prepared_count"], 1)
            self.assertEqual(decision["queue"]["awaiting_approval"], 1)
            self.assertEqual(decision["queue"]["deferred_intake"], 0)
            self.assertEqual(list((temp / "continuous-intake").glob("*.json")), [])
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])
            with database.connect() as connection:
                evidence_rows = connection.execute(
                    "SELECT status FROM private_refs WHERE kind='continuous_evidence_bundle'"
                ).fetchall()
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM applications").fetchone()[0], 2)
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 0)
            self.assertEqual([str(row["status"]) for row in evidence_rows], ["DELETED"])
            safe_output = json.dumps({"second": second, "decision": decision})
            self.assertNotIn("secure-ref:", safe_output)
            self.assertNotIn("synthetic-forward-jd-two", safe_output)

    def test_ui_deferred_retention_failure_rolls_back_the_queue_admission(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            database.set_pending_limit(1)
            fixtures = PROJECT / "tests" / "fixtures"
            service = OnboardingCenterService(PROJECT, database, onboarding)
            metadata = {
                "official_url": "https://example.com/careers/synthetic-data-analyst",
                "application_url": "https://boards.greenhouse.io/example/jobs/987654",
                "guest_available": True,
                "research_title": "Synthetic company role page",
                "evidence_excerpt": "Synthetic Data Analyst",
            }

            def local_bundle(jd_name: str) -> dict[str, tuple[str, bytes]]:
                return {
                    "jd": (".txt", (fixtures / jd_name).read_bytes()),
                    "official": (".html", (fixtures / "synthetic-greenhouse-careers.html").read_bytes()),
                    "form": (".html", (fixtures / "synthetic-material-form.html").read_bytes()),
                }

            first = service.prepare_offline_application_bundle(
                metadata=metadata, files=local_bundle("synthetic-forward-jd.txt"),
            )
            self.assertEqual(first["status"], "AWAITING_APPROVAL")
            original_import = onboarding.import_bytes

            def fail_deferred_import(kind: str, value: bytes, *, synthetic: bool = False):
                if kind == "continuous_evidence_bundle":
                    raise JobOpsError("LOCAL_RETENTION_FAILED", "Synthetic retention failure.")
                return original_import(kind, value, synthetic=synthetic)

            with mock.patch.object(onboarding, "import_bytes", side_effect=fail_deferred_import):
                with self.assertRaises(JobOpsError) as failed:
                    service.prepare_offline_application_bundle(
                        metadata=metadata, files=local_bundle("synthetic-forward-jd-two.txt"),
                    )
            self.assertEqual(failed.exception.code, "LOCAL_RETENTION_FAILED")
            status = QueueManager(database).status()
            self.assertEqual(status["awaiting_approval"], 1)
            self.assertEqual(status["deferred_intake"], 0)
            self.assertEqual(status["reserved_slots"], 0)
            self.assertEqual(list((temp / "continuous-intake").glob("*.json")), [])
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])
            with database.connect() as connection:
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM private_refs WHERE kind='continuous_evidence_bundle' AND status='ACTIVE'"
                ).fetchone()[0], 0)
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 0)

    def test_failed_preparation_releases_slot_and_deletes_new_encrypted_materials(self) -> None:
        with project_temp() as temp:
            database, onboarding, orchestrator = self.build(temp)
            self.seed_completed_context(onboarding)
            fixtures = PROJECT / "tests" / "fixtures"
            route = json.loads((fixtures / "synthetic-forward-route.json").read_text(encoding="utf-8"))
            route["research"] = {
                "title": "Example Analytics Lab update", "url": "https://example.com/news/synthetic-update",
                "source_type": "official_company", "published_at": "2026-08-12T00:00:00Z",
                "accessed_at": iso_utc(), "official": True,
                "evidence_excerpt": "Example Analytics Lab uses documented checks for synthetic dataset analysis.",
            }
            route_path = temp / "saved-route-failure.json"
            route_path.write_text(json.dumps(route), encoding="utf-8")
            original_import = onboarding.import_bytes

            def fail_after_two_generated(kind: str, value: bytes, *, synthetic: bool = False):
                if kind == "visual_evidence":
                    raise JobOpsError("LOCAL_TEST_FAILURE", "Stop after encrypted output for rollback testing.")
                return original_import(kind, value, synthetic=synthetic)

            with mock.patch.object(onboarding, "import_bytes", side_effect=fail_after_two_generated):
                with self.assertRaisesRegex(JobOpsError, "Stop after encrypted output"):
                    orchestrator.run_to_awaiting(
                        fixtures / "synthetic-forward-jd.txt",
                        profile_ref=None, master_resume_ref=None, answer_bank_ref=None,
                        route_fixture=route_path, form_fixture=fixtures / "synthetic-material-form.html",
                        research_fixture=fixtures / "synthetic-research.html", synthetic=False,
                    )
            with database.connect() as connection:
                active_generated = connection.execute(
                    "SELECT COUNT(*) FROM private_refs WHERE kind LIKE 'generated_%' AND synthetic=0 AND status='ACTIVE'"
                ).fetchone()[0]
                active_visual_or_packet = connection.execute(
                    "SELECT COUNT(*) FROM private_refs WHERE kind IN ('visual_evidence','review_packet') AND synthetic=0 AND status='ACTIVE'"
                ).fetchone()[0]
                reservation = connection.execute("SELECT status FROM queue_reservations ORDER BY created_at DESC LIMIT 1").fetchone()
            self.assertEqual(active_generated, 0)
            self.assertEqual(active_visual_or_packet, 0)
            self.assertEqual(str(reservation["status"]), "RELEASED")

    def test_manual_real_profile_tick_prepares_saved_local_evidence_and_is_idempotent(self) -> None:
        with project_temp() as temp:
            database, onboarding, orchestrator = self.build(temp)
            refs = self.seed_completed_context(onboarding)
            manager = QueueManager(database)
            manifest = {
                "schema_version": 1,
                "mode": "MANUAL_TICK_ONLY",
                "jobs": [{
                    "input": "tests/fixtures/synthetic-forward-jd.txt",
                    "profile_ref": refs["profile_ref"],
                    "master_resume_ref": refs["master_resume_ref"],
                    "answer_bank_ref": refs["answer_bank_ref"],
                    "external_claim_set_ref": refs["external_claim_set_ref"],
                    "tailoring_manifest_ref": refs["tailoring_manifest_ref"],
                    "route": "tests/fixtures/synthetic-real-offline-route.json",
                    "form": "tests/fixtures/synthetic-material-form.html",
                    "research": "tests/fixtures/synthetic-research.html",
                    "source_type": "txt",
                    "synthetic": False,
                }],
            }

            def prepare(item: dict) -> dict:
                return orchestrator.run_to_awaiting(
                    PROJECT / item["input"],
                    profile_ref=item["profile_ref"], master_resume_ref=item["master_resume_ref"],
                    answer_bank_ref=item["answer_bank_ref"],
                    external_claim_set_ref=item["external_claim_set_ref"],
                    tailoring_manifest_ref=item["tailoring_manifest_ref"],
                    route_fixture=PROJECT / item["route"], form_fixture=PROJECT / item["form"],
                    research_fixture=PROJECT / item["research"], source_type=item["source_type"],
                    synthetic=False,
                )

            first = run_continuous_intake_tick(manifest, queue_status=manager.status, prepare_job=prepare)
            self.assertEqual(first["status"], "MANUAL_TICK_COMPLETE")
            self.assertEqual(first["prepared_count"], 1)
            self.assertEqual(first["results"][0]["source_mode"], "SAVED_LOCAL_EVIDENCE")
            self.assertEqual(first["results"][0]["status"], "PREPARED")
            self.assertEqual(first["real_external_actions"], 0)
            safe_json = json.dumps(first)
            self.assertNotIn("secure-ref:", safe_json)
            self.assertNotIn("synthetic-forward-jd", safe_json)

            second = run_continuous_intake_tick(manifest, queue_status=manager.status, prepare_job=prepare)
            self.assertEqual(second["deduplicated_count"], 1)
            self.assertEqual(second["results"][0]["status"], "ALREADY_TRACKED")
            with database.connect() as connection:
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM applications").fetchone()[0], 1)
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 0)

    def test_review_decision_automatically_fills_freed_slot_from_recorded_local_batch(self) -> None:
        with project_temp() as temp:
            database, onboarding, orchestrator = self.build(temp)
            refs = self.seed_completed_context(onboarding)
            database.set_pending_limit(1)
            manager = QueueManager(database)
            store = ContinuousIntakeDescriptorStore(database, PROJECT / "schemas")

            def job(input_name: str) -> dict:
                return {
                    "input": f"tests/fixtures/{input_name}",
                    "profile_ref": refs["profile_ref"], "master_resume_ref": refs["master_resume_ref"],
                    "answer_bank_ref": refs["answer_bank_ref"],
                    "external_claim_set_ref": refs["external_claim_set_ref"],
                    "tailoring_manifest_ref": refs["tailoring_manifest_ref"],
                    "route": "tests/fixtures/synthetic-real-offline-route.json",
                    "form": "tests/fixtures/synthetic-material-form.html",
                    "research": "tests/fixtures/synthetic-research.html",
                    "source_type": "txt", "synthetic": False,
                }

            manifest = {
                "schema_version": 1, "mode": "MANUAL_TICK_ONLY",
                "jobs": [job("synthetic-forward-jd.txt"), job("synthetic-forward-jd-two.txt")],
            }

            def prepare(item: dict) -> dict:
                result = orchestrator.run_to_awaiting(
                    PROJECT / item["input"], profile_ref=item["profile_ref"],
                    master_resume_ref=item["master_resume_ref"], answer_bank_ref=item["answer_bank_ref"],
                    external_claim_set_ref=item["external_claim_set_ref"],
                    tailoring_manifest_ref=item["tailoring_manifest_ref"],
                    route_fixture=PROJECT / item["route"], form_fixture=PROJECT / item["form"],
                    research_fixture=PROJECT / item["research"], source_type="txt", synthetic=False,
                )
                if result.get("status") == "DEFERRED":
                    store.remember(str(result["intake_key"]), item)
                return result

            initial = run_continuous_intake_tick(manifest, queue_status=manager.status, prepare_job=prepare)
            self.assertEqual(initial["prepared_count"], 1)
            self.assertEqual(initial["deferred_count"], 1)
            first_application = str(initial["results"][0]["application_id"])
            service = OnboardingCenterService(PROJECT, database, onboarding)
            packet = service.review_packet(first_application)["packet"]
            decision = service.decide_review_packet({
                "application_id": first_application, "decision": "APPROVE",
                "expected_packet_hash": packet["content_hash"], "user_confirmed": True,
            })

            continued = decision["continued_intake"]
            self.assertEqual(continued["status"], "LOCAL_CONTINUATION_PROCESSED")
            self.assertEqual(continued["processed_count"], 1)
            self.assertEqual(continued["prepared_count"], 1)
            self.assertEqual(continued["results"][0]["source_mode"], "SAVED_LOCAL_EVIDENCE")
            self.assertEqual(decision["queue"]["awaiting_approval"], 1)
            self.assertEqual(decision["queue"]["deferred_intake"], 0)
            safe_json = json.dumps(continued)
            self.assertNotIn("secure-ref:", safe_json)
            self.assertNotIn("synthetic-forward-jd-two", safe_json)
            self.assertEqual(list((temp / "continuous-intake").glob("*.json")), [])
            with database.connect() as connection:
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM applications").fetchone()[0], 2)
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 0)

    def test_dynamic_questions_create_supplemental_review_without_losing_confirmed_fields(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, original_packet = self.approved_company_application(database, onboarding)
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            initial_html = (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8")
            initial_count = len(bundle["form_snapshot"]["fields"])
            prepared = service.browser_assist.prepare(
                token,
                {
                    "url": started["approved_url"],
                    "sanitized_html": initial_html,
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, initial_count + 1)],
                    "blocker_signals": [],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            live_positions = {
                f"DOM-{index:012d}": index
                for index in range(1, initial_count + 1)
            }
            self.assertEqual(
                [live_positions[item["client_ref"]] for item in prepared["fields"]],
                sorted(live_positions[item["client_ref"]] for item in prepared["fields"]),
            )
            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, _metadata = service.browser_assist.take_file(
                    token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
            field_bindings = [
                {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                for item in prepared["fields"]
            ]
            dynamic_html = initial_html.replace(
                '<button id="submit" type="submit">Submit application</button>',
                '<label for="conditional_reason">Why are you interested in this role?</label>'
                '<input id="conditional_reason" name="conditional_reason" type="text" required>'
                '<button id="submit" type="submit">Submit application</button>',
            )
            supplemental = service.browser_assist.discover_dynamic_fields(
                token,
                {
                    "url": started["approved_url"],
                    "sanitized_html": dynamic_html,
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, initial_count + 2)],
                    "blocker_signals": [],
                    "field_bindings": field_bindings,
                    "material_bindings": material_bindings,
                    "submit_events": 0,
                    "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(supplemental["status"], "SUPPLEMENTAL_REVIEW_REQUIRED")
            self.assertEqual(supplemental["dynamic_field_count"], 1)
            self.assertEqual(supplemental["real_external_actions"], 2)
            current = service.review_packet(application_id)
            self.assertEqual(current["packet_version"], int(original_packet.get("packet_version", 1)) + 1)
            self.assertEqual(current["field_resolution"]["unresolved_count"], 1)
            unresolved = current["field_resolution"]["unresolved_fields"][0]
            self.assertEqual(unresolved["answer_key"], "UNKNOWN")
            self.assertEqual(unresolved["label"], "Why are you interested in this role?")
            self.assertEqual(current["application_status"], "AWAITING_APPROVAL")

            resolved = service.resolve_application_fields({
                "application_id": application_id,
                "expected_packet_hash": current["packet"]["content_hash"],
                "resolutions": [{
                    "control_ref": unresolved["control_ref"],
                    "decision": "CONFIRMED_VALUE",
                    "value": "Synthetic finance and risk analysis role",
                }],
                "non_form_resolutions": [],
                "user_confirmed": True,
            })
            rebound = service.review_packet(application_id)
            self.assertEqual(resolved["remaining_unresolved_count"], 0)
            self.assertEqual(rebound["field_resolution"]["unresolved_count"], 0)
            approved = service.decide_review_packet({
                "application_id": application_id,
                "decision": "APPROVE",
                "expected_packet_hash": rebound["packet"]["content_hash"],
                "user_confirmed": True,
            })
            self.assertEqual(approved["status"], "APPROVED")
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 2)

    def test_successful_upload_control_replacement_does_not_create_false_dynamic_review(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            initial_html = (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8")
            initial_count = len(bundle["form_snapshot"]["fields"])
            prepared = service.browser_assist.prepare(
                token,
                {
                    "url": started["approved_url"], "sanitized_html": initial_html,
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, initial_count + 1)],
                    "blocker_signals": [],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, _ = service.browser_assist.take_file(
                    token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
            field_bindings = [
                {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                for item in prepared["fields"]
            ]
            after_upload_html = initial_html
            for snippet in (
                '        <label for="resume">Resume</label>\n        <input id="resume" name="resume" type="file" required>\n',
                '        <label for="cover_letter">Cover Letter</label>\n        <input id="cover_letter" name="cover_letter" type="file" required>\n',
                '        <label for="portfolio_file">Portfolio work sample</label>\n        <input id="portfolio_file" name="portfolio_file" type="file">\n',
            ):
                after_upload_html = after_upload_html.replace(snippet, "")
            reviewed = service.browser_assist.discover_dynamic_fields(
                token,
                {
                    "url": started["approved_url"], "sanitized_html": after_upload_html,
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, initial_count - 2)],
                    "blocker_signals": [], "field_bindings": field_bindings,
                    "material_bindings": material_bindings, "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(reviewed["status"], "NO_DYNAMIC_FIELDS")
            completed = service.browser_assist.complete(
                token,
                {
                    "field_bindings": field_bindings, "material_bindings": material_bindings,
                    "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(completed["status"], "AWAITING_USER_SUBMIT")
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 2)
            service.close()

    def test_partial_page_apply_is_audited_once_and_requires_explicit_restart(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            prepared = service.browser_assist.prepare(
                token,
                {
                    "url": started["approved_url"],
                    "sanitized_html": (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8"),
                    "client_refs": [
                        f"DOM-{index:012d}"
                        for index in range(1, len(bundle["form_snapshot"]["fields"]) + 1)
                    ],
                    "blocker_signals": [],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            first_file = prepared["files"][0]
            file_token = str(first_file["download_path"]).rsplit("/", 1)[-1]
            raw, _ = service.browser_assist.take_file(
                token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            raw[:] = b"\0" * len(raw)
            first_field = prepared["fields"][0]
            aborted = service.browser_assist.abort_page_apply(
                token,
                {
                    "cause_code": "COMPANION_CONTROL_REBIND_FAILED",
                    "failed_client_ref": first_field["client_ref"],
                    "attempted_field_bindings": [{
                        "client_ref": first_field["client_ref"],
                        "value_sha256": first_field["value_sha256"],
                    }],
                    "attempted_material_bindings": [],
                    "submit_events": 0,
                    "navigation_actions": 0,
                    "companion_tab_id": 17,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(aborted["status"], "APPLY_RESTART_REQUIRED")
            self.assertEqual(aborted["code"], "COMPANION_APPLY_RESTART_REQUIRED")
            self.assertEqual(aborted["field_attempt_count"], 1)
            self.assertEqual(aborted["file_attempt_count"], 1)
            self.assertEqual(aborted["failure_code"], "COMPANION_CONTROL_REBIND_FAILED")
            self.assertEqual(aborted["failure_page_position"], 1)
            self.assertEqual(aborted["failure_control_type"], first_field["control_type"])
            self.assertTrue(aborted["failure_field_label"])
            self.assertEqual(aborted["real_external_actions"], 2)
            self.assertFalse(aborted["submit_capability"])
            self.assertFalse(aborted["final_submit"])
            self.assertFalse(aborted["automatic_retry"])
            with database.connect() as connection:
                run = connection.execute(
                    "SELECT status FROM browser_assist_runs WHERE assist_id=?", (started["assist_id"],),
                ).fetchone()
                events = [str(row[0]) for row in connection.execute(
                    "SELECT event_type FROM browser_assist_events WHERE assist_id=? ORDER BY event_id",
                    (started["assist_id"],),
                ).fetchall()]
            self.assertEqual(str(run["status"]), "FAILED")
            self.assertEqual(events[-1], "PAGE_APPLY_ABORTED")
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 2)
            with self.assertRaises(JobOpsError) as replay:
                service.browser_assist.abort_page_apply(
                    token,
                    {
                        "cause_code": "COMPANION_CONTROL_REBIND_FAILED",
                        "attempted_field_bindings": [], "attempted_material_bindings": [],
                        "submit_events": 0, "navigation_actions": 0,
                    },
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(replay.exception.code, "BROWSER_ASSIST_TOKEN_INVALID")
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 2)
            service.close()

    def test_initial_semantic_equivalence_rebinds_ephemeral_control_refs(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            html = (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8")
            html = html.replace(">Full name</label>", ">Full name *</label>")
            prepared = service.browser_assist.prepare(
                token,
                {
                    "url": started["approved_url"], "sanitized_html": html,
                    "client_refs": [
                        f"DOM-{index:012d}"
                        for index in range(1, len(bundle["form_snapshot"]["fields"]) + 1)
                    ],
                    "blocker_signals": [],
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(prepared["status"], "LIVE_PAGE_APPROVED_FOR_ASSIST")
            self.assertGreaterEqual(prepared["field_count"], 3)
            self.assertEqual(prepared["file_count"], 3)
            service.close()

    def test_user_present_company_assist_prefills_uploads_and_stops_before_submit(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)

            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            with self.assertRaises(JobOpsError) as wrong_origin:
                service.browser_assist.pair(token, extension_origin="chrome-extension://aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa")
            self.assertEqual(wrong_origin.exception.code, "BROWSER_COMPANION_ORIGIN_FORBIDDEN")
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            field_count = len(bundle["form_snapshot"]["fields"])
            live = {
                "url": started["approved_url"],
                "sanitized_html": (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8"),
                "client_refs": [f"DOM-{index:012d}" for index in range(1, field_count + 1)],
                "blocker_signals": [],
            }
            with self.assertRaises(JobOpsError) as drifted:
                service.browser_assist.prepare(
                    token, {**live, "sanitized_html": live["sanitized_html"].replace("Full name", "Changed legal identity")},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(drifted.exception.code, "SITE_CHANGED")
            prepared = service.browser_assist.prepare(token, live, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertGreaterEqual(prepared["field_count"], 3)
            self.assertEqual(prepared["file_count"], 3)
            self.assertFalse(prepared["submit_capability"])
            self.assertTrue(prepared["stop_before_submit"])

            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, metadata = service.browser_assist.take_file(
                    token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                try:
                    self.assertEqual(sha256_bytes(bytes(raw)), item["sha256"])
                    self.assertEqual(metadata["filename"], item["filename"])
                finally:
                    raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
                with self.assertRaises(JobOpsError) as replay:
                    service.browser_assist.take_file(token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN)
                self.assertEqual(replay.exception.code, "BROWSER_FILE_TOKEN_INVALID")

            field_bindings = [
                {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                for item in prepared["fields"]
            ]
            with self.assertRaises(JobOpsError) as forbidden_submit:
                service.browser_assist.complete(
                    token,
                    {"field_bindings": field_bindings, "material_bindings": material_bindings, "submit_events": 1, "navigation_actions": 0},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(forbidden_submit.exception.code, "FINAL_SUBMIT_FORBIDDEN")
            completed = service.browser_assist.complete(
                token,
                {"field_bindings": field_bindings, "material_bindings": material_bindings, "submit_events": 0, "navigation_actions": 0},
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(completed["status"], "AWAITING_USER_SUBMIT")
            self.assertTrue(completed["user_must_click_submit"])
            self.assertFalse(completed["submit_capability"])

            observed = service.browser_assist.submit_observed(
                token,
                {"url": started["approved_url"], "trusted_user_event": True, "event_hash": "sha256:" + "1" * 64},
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(observed["submit_performed_by"], "USER")
            confirmed = service.browser_assist.observe_result(
                token,
                {
                    "url": "https://example.com/careers/apply/synthetic-data-analyst",
                    "success_markers": ["APPLICATION_RECEIVED"], "failure_markers": [],
                    "invalid_control_count": 0, "form_present": False,
                    "submit_control_present": False, "success_route": True,
                    "page_fingerprint": "sha256:" + "2" * 64,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(confirmed["status"], "CONFIRMED")
            audit = audit_real_external_actions(database)
            self.assertEqual(audit["attempt_count"], 3)
            self.assertEqual(audit["real_external_actions"], 2)
            with database.connect() as connection:
                actions = [
                    (str(row["action"]), int(row["real_side_effect"]))
                    for row in connection.execute("SELECT action,real_side_effect FROM external_action_session_uses ORDER BY used_at")
                ]
                self.assertEqual(connection.execute("SELECT status FROM applications WHERE application_id=?", (application_id,)).fetchone()[0], "CONFIRMED")
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM receipts WHERE application_id=? AND source='browser-companion' AND verified=1", (application_id,)).fetchone()[0], 1)
            self.assertEqual(actions, [
                ("inspect_application_form", 0),
                ("prefill_application_form", 1),
                ("upload_materials", 1),
            ])
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])

    def test_local_ai_agent_uses_same_approved_page_tools_and_cannot_submit(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)

            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            paired = service.pair_local_agent_assist({"assist_token": token, "user_confirmed": True})
            self.assertEqual(paired["status"], "LOCAL_AI_AGENT_PAIRED")
            self.assertEqual(paired["execution_channel"], "LOCAL_AI_AGENT")
            self.assertEqual(paired["model_private_values"], 0)

            bundle, _, _ = service.browser_assist._bundle_manager.load_current(application_id)
            field_count = len(bundle["form_snapshot"]["fields"])
            prepared = service.prepare_local_agent_assist({
                "assist_token": token,
                "page": {
                    "url": started["approved_url"],
                    "sanitized_html": (PROJECT / "tests" / "fixtures" / "synthetic-material-form.html").read_text(encoding="utf-8"),
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, field_count + 1)],
                    "blocker_signals": [],
                },
            })
            self.assertEqual(prepared["status"], "LIVE_PAGE_APPROVED_FOR_ASSIST")
            self.assertEqual(prepared["execution_channel"], "LOCAL_AI_AGENT")
            self.assertEqual(prepared["model_private_values"], 0)
            self.assertFalse(prepared["submit_capability"])
            self.assertTrue(prepared["stop_before_submit"])

            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, metadata = service.take_local_agent_assist_file(
                    assist_token=token, file_token=file_token,
                )
                try:
                    self.assertEqual(sha256_bytes(bytes(raw)), item["sha256"])
                    self.assertEqual(metadata["filename"], item["filename"])
                finally:
                    raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
            field_bindings = [
                {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                for item in prepared["fields"]
            ]
            with self.assertRaises(JobOpsError) as forbidden:
                service.complete_local_agent_assist({
                    "assist_token": token,
                    "evidence": {
                        "field_bindings": field_bindings,
                        "material_bindings": material_bindings,
                        "submit_events": 1,
                        "navigation_actions": 0,
                    },
                })
            self.assertEqual(forbidden.exception.code, "FINAL_SUBMIT_FORBIDDEN")
            completed = service.complete_local_agent_assist({
                "assist_token": token,
                "evidence": {
                    "field_bindings": field_bindings,
                    "material_bindings": material_bindings,
                    "submit_events": 0,
                    "navigation_actions": 0,
                },
            })
            self.assertEqual(completed["status"], "AWAITING_USER_SUBMIT")
            self.assertEqual(completed["execution_channel"], "LOCAL_AI_AGENT")
            self.assertEqual(completed["final_submit"], "USER_ONLY")
            self.assertFalse(completed["submit_capability"])
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 2)
            service.browser_assist.stop(user_confirmed=True)

    def test_expired_unchanged_packet_is_renewed_only_by_fresh_start_confirmation(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            with database.connect() as connection:
                connection.execute(
                    "UPDATE approvals SET expires_at='2000-01-01T00:00:00Z' WHERE application_id=? AND status='APPROVED'",
                    (application_id,),
                )
            with self.assertRaises(JobOpsError) as confirmation:
                service.start_browser_assist({"application_id": application_id, "user_confirmed": False})
            self.assertEqual(confirmation.exception.code, "EXPLICIT_CONFIRMATION_REQUIRED")

            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            self.assertTrue(started["approval_renewed"])
            with database.connect() as connection:
                rows = connection.execute(
                    "SELECT status FROM approvals WHERE application_id=? ORDER BY issued_at",
                    (application_id,),
                ).fetchall()
                event = connection.execute(
                    "SELECT event_type,from_state,to_state FROM events WHERE application_id=? AND event_type='APPROVAL_RENEWED' ORDER BY event_id DESC LIMIT 1",
                    (application_id,),
                ).fetchone()
            self.assertEqual([str(row["status"]) for row in rows], ["INVALIDATED", "APPROVED"])
            self.assertEqual((event["event_type"], event["from_state"], event["to_state"]), (
                "APPROVAL_RENEWED", "APPROVED", "APPROVED",
            ))
            service.browser_assist.stop(user_confirmed=True)

    def test_unreadable_result_requires_manual_answer_and_never_retries(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            token, prepared = self.prepared_browser_assist(service, application_id)
            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, _ = service.browser_assist.take_file(token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN)
                raw[:] = b"\0" * len(raw)
                material_bindings.append({"client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"]})
            service.browser_assist.complete(
                token,
                {
                    "field_bindings": [{"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]} for item in prepared["fields"]],
                    "material_bindings": material_bindings, "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            started = next(item for item in service.browser_assist._leases.values() if item.token == token)
            service.browser_assist.submit_observed(
                token,
                {"url": str(started.source_route["current_url"]), "trusted_user_event": True, "event_hash": "sha256:" + "3" * 64},
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            unknown = service.browser_assist.result_unavailable(
                token, {"reason": "PAGE_UNAVAILABLE"}, extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(unknown["status"], "SUBMISSION_UNKNOWN")
            self.assertEqual(unknown["question"], "是否提交成功？")
            self.assertFalse(unknown["automatic_retry"])
            resolved = service.resolve_browser_assist_unknown({
                "application_id": application_id, "submitted": False, "user_confirmed": True,
            })
            self.assertEqual(resolved["status"], "AWAITING_APPROVAL")
            self.assertFalse(resolved["automatic_retry"])
            with database.connect() as connection:
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0], 3)
                self.assertEqual(connection.execute("SELECT COUNT(*) FROM browser_assist_runs").fetchone()[0], 1)
                self.assertEqual(connection.execute("SELECT status FROM applications WHERE application_id=?", (application_id,)).fetchone()[0], "AWAITING_APPROVAL")

    def test_process_close_during_user_submit_window_becomes_unknown_without_retry(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            token, prepared = self.prepared_browser_assist(service, application_id)
            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, _ = service.browser_assist.take_file(
                    token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
            service.browser_assist.complete(
                token,
                {
                    "field_bindings": [
                        {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                        for item in prepared["fields"]
                    ],
                    "material_bindings": material_bindings,
                    "submit_events": 0,
                    "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            with database.connect() as connection:
                attempts_before_close = int(connection.execute(
                    "SELECT COUNT(*) FROM external_action_attempts"
                ).fetchone()[0])
            service.close()
            with database.connect() as connection:
                run = connection.execute(
                    "SELECT status FROM browser_assist_runs WHERE application_id=?",
                    (application_id,),
                ).fetchone()
                application = connection.execute(
                    "SELECT status FROM applications WHERE application_id=?", (application_id,),
                ).fetchone()
                attempts_after_close = int(connection.execute(
                    "SELECT COUNT(*) FROM external_action_attempts"
                ).fetchone()[0])
            self.assertEqual(str(run["status"]), "SUBMISSION_UNKNOWN")
            self.assertEqual(str(application["status"]), "SUBMISSION_UNKNOWN")
            self.assertEqual(attempts_after_close, attempts_before_close)

    def test_v2_workday_multi_page_handoff_navigation_upload_and_human_submit(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, packet = self.approved_workday_v2_application(database, onboarding)
            self.assertEqual(packet["source_route"]["provider"], "workday")
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            self.assertEqual(started["protocol_version"], 2)
            self.assertTrue(started["multi_page"])
            token = str(started["assist_token"])
            paired = service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(paired["provider"], "workday")
            fixtures = PROJECT / "tests" / "fixtures"

            def live(
                name: str, count: int, *, signals: list[str] | None = None,
                url: str | None = None, document: str = "A" * 32,
            ) -> dict:
                return {
                    "url": url or started["approved_url"],
                    "sanitized_html": (fixtures / name).read_text(encoding="utf-8"),
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, count + 1)],
                    "companion_tab_id": 42,
                    "document_instance_id": "DOC-" + document,
                    "blocker_signals": list(signals or []),
                }

            first = service.browser_assist.prepare(
                token, live("synthetic-v2-workday-step-1.html", 2),
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(first["page_kind"], "INTERMEDIATE")
            self.assertEqual(first["field_count"], 1)
            self.assertEqual(first["manual_field_count"], 0)
            first_done = service.browser_assist.complete(
                token,
                {
                    "field_bindings": [{"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]} for item in first["fields"]],
                    "material_bindings": [], "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(first_done["status"], "MANUAL_NAVIGATION_REQUIRED")
            self.assertFalse(first_done["manual_navigation"]["programmatic_allowed"])
            first_challenge = first_done["manual_navigation"]["challenge"]
            self.assertEqual(
                (parse_iso(first_challenge["expires_at"]) - parse_iso(first_challenge["issued_at"])).total_seconds(),
                15 * 60,
            )
            self.assertEqual(first["navigation"]["control_type"], "submit")
            self.assertIsNone(first["navigation"]["authorization_token"])
            with self.assertRaises(JobOpsError) as manual_only:
                service.browser_assist.authorize_navigation(
                    token,
                    {"client_ref": first["navigation"]["client_ref"], "authorization_token": "", "form_valid": True, "submit_events": 0},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(manual_only.exception.code, "NAVIGATION_REQUIRES_USER_CLICK")
            first_manual = manual_navigation_evidence(service, token, first_done)
            step_two_url = started["approved_url"] + "/apply/step-2"
            resumed_first = service.browser_assist.resume_manual_navigation(
                token, {
                    **live("synthetic-v2-workday-step-2.html", 4, url=step_two_url, document="B" * 32),
                    **first_manual,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(resumed_first["navigation_performed_by"], "USER")
            self.assertEqual(resumed_first["current_step"], 2)

            handoff = service.browser_assist.prepare(
                token, live("synthetic-v2-login.html", 3, signals=["LOGIN"], url=step_two_url, document="B" * 32),
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(handoff["status"], "HANDOFF_REQUIRED")
            self.assertEqual(handoff["handoff_kind"], "LOGIN")
            self.assertFalse(handoff["account_creation_capability"])
            second = service.browser_assist.prepare(
                token, live("synthetic-v2-workday-step-2.html", 4, url=step_two_url, document="B" * 32),
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(second["current_step"], 2)
            self.assertEqual(second["field_count"], 1)
            self.assertEqual(second["manual_field_count"], 1)
            second_done = service.browser_assist.complete(
                token,
                {
                    "field_bindings": [{"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]} for item in second["fields"]],
                    "material_bindings": [], "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(second_done["status"], "MANUAL_NAVIGATION_REQUIRED")
            second_manual = manual_navigation_evidence(service, token, second_done)
            step_three_url = started["approved_url"] + "/apply/step-3"
            service.browser_assist.resume_manual_navigation(
                token, {
                    **live("synthetic-v2-workday-step-3.html", 5, url=step_three_url, document="C" * 32),
                    **second_manual,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )

            final = service.browser_assist.prepare(
                token, live("synthetic-v2-workday-step-3.html", 5, url=step_three_url, document="C" * 32),
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(final["page_kind"], "FINAL_REVIEW")
            self.assertEqual(final["file_count"], 3)
            self.assertEqual(len(final["final_submit_client_refs"]), 1)
            material_bindings = []
            for item in final["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, _ = service.browser_assist.take_file(
                    token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
            final_done = service.browser_assist.complete(
                token,
                {"field_bindings": [], "material_bindings": material_bindings, "submit_events": 0, "navigation_actions": 0},
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(final_done["status"], "AWAITING_USER_SUBMIT")
            service.browser_assist.submit_observed(
                token,
                {"url": started["approved_url"], "trusted_user_event": True, "event_hash": "sha256:" + "6" * 64},
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            confirmed = service.browser_assist.observe_result(
                token,
                {
                    "url": started["approved_url"], "success_markers": ["APPLICATION_RECEIVED"],
                    "failure_markers": [], "invalid_control_count": 0, "form_present": False,
                    "submit_control_present": False, "success_route": True,
                    "page_fingerprint": "sha256:" + "7" * 64,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(confirmed["status"], "CONFIRMED")
            audit = audit_real_external_actions(database)
            self.assertEqual(audit["attempt_count"], 7)
            self.assertEqual(audit["real_external_actions"], 3)
            with database.connect() as connection:
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM external_action_session_uses WHERE action='navigate_application_step'"
                ).fetchone()[0], 0)
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM browser_assist_events WHERE event_type='MANUAL_NEXT_PAGE_OBSERVED'"
                ).fetchone()[0], 2)
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM browser_assist_events WHERE event_type='USER_HANDOFF_REQUIRED'"
                ).fetchone()[0], 1)
                serialized = " ".join(str(row[0]) for row in connection.execute(
                    "SELECT payload_json FROM events WHERE application_id=?", (application_id,),
                ))
            self.assertNotIn("Additional information", serialized)
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])

    def test_v2_handoffs_cross_origin_and_repeated_page_fail_closed(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_workday_v2_application(database, onboarding)
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            fixture = (PROJECT / "tests" / "fixtures" / "synthetic-v2-workday-step-1.html").read_text(encoding="utf-8")

            def page(*, url: str | None = None, signals: list[str] | None = None) -> dict:
                return {
                    "url": url or started["approved_url"],
                    "sanitized_html": fixture,
                    "client_refs": ["DOM-000000000001", "DOM-000000000002"],
                    "companion_tab_id": 42,
                    "document_instance_id": "DOC-" + "D" * 32,
                    "blocker_signals": list(signals or []),
                }

            with self.assertRaises(JobOpsError) as spoofed_handoff:
                service.browser_assist.prepare(
                    token, page(url="https://evil.example.test/apply", signals=["CAPTCHA"]),
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(spoofed_handoff.exception.code, "FORM_ROUTE_BINDING_CHANGED")

            captcha = service.browser_assist.prepare(
                token, page(signals=["CAPTCHA"]), extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(captcha["status"], "HANDOFF_REQUIRED")
            self.assertEqual(captcha["handoff_kind"], "CAPTCHA")
            self.assertFalse(captcha["captcha_bypass_capability"])
            account = service.browser_assist.prepare(
                token, page(signals=["ACCOUNT_CREATION"]), extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(account["status"], "HANDOFF_REQUIRED")
            self.assertEqual(account["handoff_kind"], "ACCOUNT_OR_LOGIN")
            self.assertTrue(account["existing_account_only"])
            self.assertFalse(account["account_creation_capability"])

            with self.assertRaises(JobOpsError) as wrong_origin:
                service.browser_assist.prepare(
                    token, page(url="https://evil.example.test/apply"),
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(wrong_origin.exception.code, "FORM_ROUTE_BINDING_CHANGED")

            prepared = service.browser_assist.prepare(token, page(), extension_origin=COMPANION_EXTENSION_ORIGIN)
            completed = service.browser_assist.complete(
                token,
                {
                    "field_bindings": [
                        {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                        for item in prepared["fields"]
                    ],
                    "material_bindings": [], "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            with self.assertRaises(JobOpsError) as repeated:
                service.browser_assist.resume_manual_navigation(
                    token,
                    {
                        **page(),
                        **manual_navigation_evidence(service, token, completed),
                    },
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(repeated.exception.code, "NAVIGATION_DID_NOT_ADVANCE")
            with self.assertRaises(JobOpsError) as replayed:
                service.browser_assist.resume_manual_navigation(
                    token,
                    {**page(), **manual_navigation_evidence(service, token, completed)},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(replayed.exception.code, "MANUAL_NAVIGATION_CHALLENGE_REPLAYED")
            self.assertEqual(completed["status"], "MANUAL_NAVIGATION_REQUIRED")
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 1)
            self.assertEqual(list((onboarding.store.private_root / "staging").iterdir()), [])
            service.browser_assist.stop(user_confirmed=True)
            with database.connect() as connection:
                self.assertEqual(connection.execute(
                    "SELECT COUNT(*) FROM external_action_sessions WHERE status='AUTHORIZED'"
                ).fetchone()[0], 0)

    def test_manual_navigation_challenge_rejects_forgery_and_different_job_identity(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_workday_v2_application(database, onboarding)
            started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            token = str(started["assist_token"])
            service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            fixtures = PROJECT / "tests" / "fixtures"

            def page(name: str, count: int, *, url: str | None = None, document: str = "E" * 32) -> dict:
                return {
                    "url": url or started["approved_url"],
                    "sanitized_html": (fixtures / name).read_text(encoding="utf-8"),
                    "client_refs": [f"DOM-{index:012d}" for index in range(1, count + 1)],
                    "companion_tab_id": 52,
                    "document_instance_id": "DOC-" + document,
                    "blocker_signals": [],
                }

            prepared = service.browser_assist.prepare(
                token, page("synthetic-v2-workday-step-1.html", 2),
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            completed = service.browser_assist.complete(
                token,
                {
                    "field_bindings": [
                        {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                        for item in prepared["fields"]
                    ],
                    "material_bindings": [], "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            challenge = completed["manual_navigation"]["challenge"]
            self.assertEqual(challenge["stage"], "MANUAL_NAVIGATION_REQUIRED")
            self.assertEqual(challenge["tab_id"], 52)
            self.assertEqual(
                (parse_iso(challenge["expires_at"]) - parse_iso(challenge["issued_at"])).total_seconds(),
                15 * 60,
            )
            evidence = manual_navigation_evidence(service, token, completed)
            cancelled = {**evidence, "manual_navigation_default_prevented": True}
            with self.assertRaises(JobOpsError) as prevented:
                service.browser_assist.resume_manual_navigation(
                    token,
                    {**page("synthetic-v2-workday-step-2.html", 4, document="F" * 32), **cancelled},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(prevented.exception.code, "MANUAL_NAVIGATION_EVENT_CANCELLED")
            self.assertEqual(
                service.browser_assist._leases[token].manual_challenge_id,
                challenge["challenge_id"],
            )
            forged = {**evidence, "event_hash": "sha256:" + "0" * 64}
            with self.assertRaises(JobOpsError) as arbitrary_hash:
                service.browser_assist.resume_manual_navigation(
                    token,
                    {**page("synthetic-v2-workday-step-2.html", 4, document="F" * 32), **forged},
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(arbitrary_hash.exception.code, "MANUAL_NAVIGATION_EVIDENCE_INVALID")

            different_job_url = started["approved_url"].replace("/job/123", "/job/999") + "/apply/step-2"
            with self.assertRaises(JobOpsError) as changed_identity:
                service.browser_assist.resume_manual_navigation(
                    token,
                    {
                        **page(
                            "synthetic-v2-workday-step-2.html", 4,
                            url=different_job_url, document="F" * 32,
                        ),
                        **evidence,
                    },
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(changed_identity.exception.code, "FORM_ROUTE_IDENTITY_CHANGED")
            with self.assertRaises(JobOpsError) as replayed:
                service.browser_assist.resume_manual_navigation(
                    token,
                    {
                        **page("synthetic-v2-workday-step-2.html", 4, document="F" * 32),
                        **evidence,
                    },
                    extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
            self.assertEqual(replayed.exception.code, "MANUAL_NAVIGATION_CHALLENGE_REPLAYED")
            self.assertEqual(audit_real_external_actions(database)["real_external_actions"], 1)
            service.browser_assist.stop(user_confirmed=True)

    def test_browser_companion_reload_revokes_pre_submit_checkpoints_and_allows_new_lease(self) -> None:
        fixtures = PROJECT / "tests" / "fixtures"

        for target_status in ("PAGE_REVIEW_REQUIRED", "HANDOFF_REQUIRED", "AWAITING_NAVIGATION"):
            with self.subTest(target_status=target_status), project_temp() as temp:
                database, onboarding, _ = self.build(temp)
                self.seed_completed_context(onboarding)
                service, application_id, _ = self.approved_workday_v2_application(
                    database, onboarding, initial_form="synthetic-v2-workday-step-1-explicit-button.html",
                )
                started = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
                token = str(started["assist_token"])
                service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
                lease = service.browser_assist._leases[token]
                old_assist_id = lease.assist_id
                old_session_id = lease.session_id
                live = {
                    "url": started["approved_url"],
                    "sanitized_html": (fixtures / "synthetic-v2-workday-step-1-explicit-button.html").read_text(encoding="utf-8"),
                    "client_refs": ["DOM-000000000001", "DOM-000000000002"],
                    "blocker_signals": [],
                }

                if target_status == "HANDOFF_REQUIRED":
                    handoff = service.browser_assist.prepare(
                        token, {**live, "blocker_signals": ["LOGIN"]},
                        extension_origin=COMPANION_EXTENSION_ORIGIN,
                    )
                    self.assertEqual(handoff["status"], target_status)
                else:
                    prepared = service.browser_assist.prepare(
                        token, live, extension_origin=COMPANION_EXTENSION_ORIGIN,
                    )
                    completed = service.browser_assist.complete(
                        token,
                        {
                            "field_bindings": [
                                {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                                for item in prepared["fields"]
                            ],
                            "material_bindings": [], "submit_events": 0, "navigation_actions": 0,
                        },
                        extension_origin=COMPANION_EXTENSION_ORIGIN,
                    )
                    self.assertEqual(completed["status"], "PAGE_REVIEW_REQUIRED")
                    if target_status == "AWAITING_NAVIGATION":
                        with self.assertRaises(JobOpsError) as stale_proof:
                            service.browser_assist.authorize_navigation(
                                token,
                                {
                                    "client_ref": prepared["navigation"]["client_ref"],
                                    "authorization_token": prepared["navigation"]["authorization_token"],
                                    "form_valid": True, "submit_events": 0,
                                    "page_content_hash": "sha256:" + "0" * 64,
                                    "control_semantics_hash": prepared["navigation"]["control_semantics_hash"],
                                },
                                extension_origin=COMPANION_EXTENSION_ORIGIN,
                            )
                        self.assertEqual(stale_proof.exception.code, "NAVIGATION_AUTHORIZATION_INVALID")
                        authorized = service.browser_assist.authorize_navigation(
                            token,
                            {
                                "client_ref": prepared["navigation"]["client_ref"],
                                "authorization_token": prepared["navigation"]["authorization_token"],
                                "form_valid": True, "submit_events": 0,
                                "page_content_hash": prepared["navigation"]["page_content_hash"],
                                "control_semantics_hash": prepared["navigation"]["control_semantics_hash"],
                            },
                            extension_origin=COMPANION_EXTENSION_ORIGIN,
                        )
                        self.assertEqual(authorized["status"], "NAVIGATION_AUTHORIZED")

                self.assertEqual(lease.status, target_status)
                if target_status == "PAGE_REVIEW_REQUIRED":
                    self.assertIsNotNone(lease.navigation_token)
                with database.connect() as connection:
                    counts_before_reload = (
                        int(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0]),
                        int(connection.execute("SELECT COUNT(*) FROM external_action_session_uses").fetchone()[0]),
                        int(connection.execute(
                            "SELECT COUNT(*) FROM external_action_session_uses WHERE action='navigate_application_step'"
                        ).fetchone()[0]),
                        int(connection.execute(
                            "SELECT COUNT(*) FROM browser_assist_events WHERE event_type='NEXT_PAGE_OBSERVED'"
                        ).fetchone()[0]),
                    )

                with self.assertRaises(JobOpsError) as reloaded:
                    service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
                self.assertEqual(reloaded.exception.code, "BROWSER_ASSIST_RESTART_REQUIRED")
                self.assertEqual(reloaded.exception.details["prior_status"], target_status)
                self.assertFalse(reloaded.exception.details["automatic_retry"])
                self.assertEqual(lease.status, "REVOKED")
                self.assertIsNone(lease.navigation_ref)
                self.assertIsNone(lease.navigation_token)
                self.assertNotIn(token, service.browser_assist._leases)

                with database.connect() as connection:
                    self.assertEqual(connection.execute(
                        "SELECT status FROM browser_assist_runs WHERE assist_id=?", (old_assist_id,),
                    ).fetchone()[0], "REVOKED")
                    self.assertEqual(connection.execute(
                        "SELECT status FROM external_action_sessions WHERE session_id=?", (old_session_id,),
                    ).fetchone()[0], "REVOKED")
                    self.assertEqual(connection.execute(
                        "SELECT status FROM applications WHERE application_id=?", (application_id,),
                    ).fetchone()[0], "APPROVED")
                    counts_after_reload = (
                        int(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0]),
                        int(connection.execute("SELECT COUNT(*) FROM external_action_session_uses").fetchone()[0]),
                        int(connection.execute(
                            "SELECT COUNT(*) FROM external_action_session_uses WHERE action='navigate_application_step'"
                        ).fetchone()[0]),
                        int(connection.execute(
                            "SELECT COUNT(*) FROM browser_assist_events WHERE event_type='NEXT_PAGE_OBSERVED'"
                        ).fetchone()[0]),
                    )
                self.assertEqual(counts_after_reload, counts_before_reload)

                fresh = service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
                self.assertNotEqual(fresh["assist_id"], old_assist_id)
                self.assertNotEqual(fresh["assist_token"], token)
                self.assertFalse(fresh.get("resumed", False))
                with database.connect() as connection:
                    self.assertEqual(connection.execute(
                        "SELECT status FROM browser_assist_runs WHERE assist_id=?", (fresh["assist_id"],),
                    ).fetchone()[0], "PAIRING")
                    self.assertEqual(connection.execute(
                        "SELECT status FROM external_action_sessions WHERE session_id=(SELECT session_id FROM browser_assist_runs WHERE assist_id=?)",
                        (fresh["assist_id"],),
                    ).fetchone()[0], "AUTHORIZED")
                    self.assertEqual(
                        int(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0]),
                        counts_before_reload[0],
                    )
                service.browser_assist.stop(user_confirmed=True)

    def test_browser_companion_reload_during_user_submit_window_becomes_unknown_and_never_retries(self) -> None:
        with project_temp() as temp:
            database, onboarding, _ = self.build(temp)
            self.seed_completed_context(onboarding)
            service, application_id, _ = self.approved_company_application(database, onboarding)
            token, prepared = self.prepared_browser_assist(service, application_id)
            material_bindings = []
            for item in prepared["files"]:
                file_token = str(item["download_path"]).rsplit("/", 1)[-1]
                raw, _ = service.browser_assist.take_file(
                    token, file_token, extension_origin=COMPANION_EXTENSION_ORIGIN,
                )
                raw[:] = b"\0" * len(raw)
                material_bindings.append({
                    "client_ref": item["client_ref"], "purpose": item["purpose"], "sha256": item["sha256"],
                })
            completed = service.browser_assist.complete(
                token,
                {
                    "field_bindings": [
                        {"client_ref": item["client_ref"], "value_sha256": item["value_sha256"]}
                        for item in prepared["fields"]
                    ],
                    "material_bindings": material_bindings, "submit_events": 0, "navigation_actions": 0,
                },
                extension_origin=COMPANION_EXTENSION_ORIGIN,
            )
            self.assertEqual(completed["status"], "AWAITING_USER_SUBMIT")
            lease = service.browser_assist._leases[token]
            old_assist_id = lease.assist_id
            with database.connect() as connection:
                actions_before_reload = int(connection.execute(
                    "SELECT COUNT(*) FROM external_action_attempts"
                ).fetchone()[0])

            with self.assertRaises(JobOpsError) as reloaded:
                service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(reloaded.exception.code, "BROWSER_ASSIST_SUBMISSION_UNKNOWN")
            self.assertFalse(reloaded.exception.details["automatic_retry"])
            self.assertEqual(lease.status, "SUBMISSION_UNKNOWN")
            self.assertNotIn(token, service.browser_assist._leases)
            with database.connect() as connection:
                self.assertEqual(connection.execute(
                    "SELECT status FROM browser_assist_runs WHERE assist_id=?", (old_assist_id,),
                ).fetchone()[0], "SUBMISSION_UNKNOWN")
                self.assertEqual(connection.execute(
                    "SELECT status FROM applications WHERE application_id=?", (application_id,),
                ).fetchone()[0], "SUBMISSION_UNKNOWN")
                self.assertEqual(connection.execute(
                    "SELECT status FROM approvals WHERE application_id=? ORDER BY issued_at DESC LIMIT 1", (application_id,),
                ).fetchone()[0], "CONSUMED")
                event_payload = json.loads(connection.execute(
                    "SELECT payload_json FROM events WHERE application_id=? AND event_type='SUBMISSION_EVIDENCE_UNKNOWN' ORDER BY event_id DESC LIMIT 1",
                    (application_id,),
                ).fetchone()[0])
                self.assertFalse(event_payload["automatic_retry"])
                self.assertEqual(
                    int(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0]),
                    actions_before_reload,
                )

            with self.assertRaises(JobOpsError) as restart:
                service.start_browser_assist({"application_id": application_id, "user_confirmed": True})
            self.assertEqual(restart.exception.code, "APPLICATION_NOT_APPROVED")
            with self.assertRaises(JobOpsError) as stale_token:
                service.browser_assist.pair(token, extension_origin=COMPANION_EXTENSION_ORIGIN)
            self.assertEqual(stale_token.exception.code, "BROWSER_ASSIST_TOKEN_INVALID")
            with database.connect() as connection:
                self.assertEqual(
                    int(connection.execute("SELECT COUNT(*) FROM external_action_attempts").fetchone()[0]),
                    actions_before_reload,
                )


if __name__ == "__main__":
    unittest.main()
